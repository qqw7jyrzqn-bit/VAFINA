from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.views.generic import ListView, DetailView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from django.db.models import Q, Count, Avg, F, ExpressionWrapper, DurationField
from django.db.models.functions import TruncDate, TruncWeek
from django.utils import timezone
from django.http import HttpResponse, JsonResponse, FileResponse
from django.db import IntegrityError, transaction
from datetime import datetime, timedelta, date
import json
import os
import re
import time
import random
import uuid

from .models import Document, DocumentTemplate, DocumentHistory, Notification, WorkflowStep, ChatMessage
from .permissions import (
    role_required, clerk_required, manager_required, clerk_or_manager_required,
    RoleRequiredMixin, ClerkRequiredMixin, ManagerRequiredMixin, ClerkOrManagerRequiredMixin,
    can_edit_document, can_delete_document, can_approve_document, can_manage_templates, can_view_all_documents,
    can_create_document
)
from .email_utils import (
    send_notification_email, send_document_assigned_email, send_document_approved_email,
    send_document_rejected_email, send_workflow_step_notification, send_chat_message_email
)
from .forms import (DocumentForm, DocumentTemplateForm, DocumentFromTemplateForm, 
                    WorkflowRouteForm, DocumentFilterForm, ApprovalForm)
from accounts.models import UserProfile
from .office_utils import generate_document_from_template


def generate_unique_registry_number():
    """Генерирует уникальный регистрационный номер с защитой от дубликатов"""
    now = timezone.now()
    year = now.year
    month = now.strftime('%m')
    
    max_attempts = 50
    for attempt in range(max_attempts):
        try:
            with transaction.atomic():
                # Блокируем последний документ для чтения
                last_doc = Document.objects.filter(
                    registry_number__startswith=f"{year}/{month}/"
                ).select_for_update().order_by('-registry_number').first()
                
                if last_doc and last_doc.registry_number:
                    try:
                        last_number = int(last_doc.registry_number.split('/')[-1])
                        new_number = last_number + 1
                    except (ValueError, IndexError):
                        new_number = 1
                else:
                    new_number = 1
                
                # Для первых 30 попыток используем последовательные номера
                if attempt < 30:
                    registry_number = f"{year}/{month}/{new_number:04d}"
                # Для попыток 30-45 используем timestamp
                elif attempt < 45:
                    timestamp = (int(time.time() * 1000) % 10000) + attempt
                    registry_number = f"{year}/{month}/{timestamp:04d}"
                # Последние 5 попыток - UUID
                else:
                    unique_suffix = str(uuid.uuid4().hex)[:4].upper()
                    registry_number = f"{year}/{month}/{unique_suffix}"
                
                # Проверяем уникальность
                if not Document.objects.filter(registry_number=registry_number).exists():
                    return registry_number
                
        except Exception:
            pass
        
        # Небольшая задержка перед повторной попыткой
        time.sleep(random.uniform(0.001, 0.005))
    
    # Абсолютный fallback - микросекунды
    timestamp = int(time.time() * 1000000) % 100000000
    return f"{year}/{month}/{timestamp}"
@login_required
def dashboard(request):
    """Главная страница - аналитическая панель"""
    from datetime import timedelta
    from django.db.models import Q, Count, Avg
    from .models import Task, DocumentComment, AuditLog
    
    user = request.user
    profile = getattr(user, 'profile', None)
    
    # Определяем видимые документы в зависимости от роли
    if profile and profile.role in ['admin', 'clerk']:
        # Администратор и делопроизводитель видят всё
        visible_docs = Document.objects.all()
        can_see_all = True
    elif profile and profile.role == 'manager':
        # Руководитель видит все документы
        visible_docs = Document.objects.all()
        can_see_all = True
    else:
        # Сотрудник видит только свои документы
        visible_docs = Document.objects.filter(
            Q(created_by=user) | Q(assigned_to=user)
        )
        can_see_all = False
    
    # Основная статистика
    total_documents = visible_docs.count()
    drafts = visible_docs.filter(status='draft').count()
    sent_for_approval = visible_docs.filter(status='sent_for_approval').count()
    coordination = visible_docs.filter(status='coordination').count()
    approval = visible_docs.filter(status='approval').count()
    approved = visible_docs.filter(status='approved').count()
    execution = visible_docs.filter(status='execution').count()
    rejected = visible_docs.filter(status='rejected').count()
    returned = visible_docs.filter(status='returned').count()
    archived = visible_docs.filter(status='archived').count()
    
    # Документы на согласовании (все кроме черновика, утвержденных и архивных)
    in_approval_statuses = ['sent_for_approval', 'coordination', 'approval']
    in_approval_count = visible_docs.filter(status__in=in_approval_statuses).count()
    
    # Просроченные документы
    today = timezone.now().date()
    overdue = visible_docs.filter(
        deadline__lt=today,
        status__in=['draft', 'sent_for_approval', 'coordination', 'approval', 'execution']
    ).count()
    
    # Документы с приближающимся дедлайном (7 дней)
    upcoming_deadline = today + timedelta(days=7)
    urgent = visible_docs.filter(
        deadline__range=[today, upcoming_deadline],
        status__in=['draft', 'sent_for_approval', 'coordination', 'approval', 'execution']
    ).count()
    
    # Недавние документы  
    recent_documents = visible_docs.select_related('created_by', 'assigned_to', 'template').order_by('-created_at')[:10]
    
    # Статистика по типам документов
    documents_by_type = visible_docs.filter(template__isnull=False).values(
        'template__type'
    ).annotate(count=Count('id')).order_by('-count')
    
    # Статистика по статусам
    documents_by_status = visible_docs.values('status').annotate(count=Count('id'))
    
    # Статистика по отделам (для админа и делопроизводителя)
    documents_by_department = []
    if can_see_all:
        documents_by_department = visible_docs.filter(
            created_by__profile__isnull=False
        ).values(
            'created_by__profile__department'
        ).annotate(count=Count('id')).order_by('-count')[:10]
    
    # Активность за последние 30 дней
    thirty_days_ago = timezone.now() - timedelta(days=30)
    recent_activity = visible_docs.filter(created_at__gte=thirty_days_ago).count()
    
    # Задачи пользователя
    user_tasks = Task.objects.filter(assigned_to=user).exclude(status='completed')
    pending_tasks = user_tasks.filter(status='pending').count()
    inprogress_tasks = user_tasks.filter(status='in_progress').count()
    overdue_tasks = user_tasks.filter(
        deadline__lt=timezone.now(),
        status__in=['pending', 'in_progress']
    ).count()
    
    # Непрочитанные уведомления
    unread_notifications = user.notifications.filter(is_read=False).count()
    
    # Подписанные документы
    signed_documents = visible_docs.filter(is_signed=True).count()
    
    # Статистика по версиям (средняя версия документов)
    avg_version = visible_docs.aggregate(avg=Avg('version'))['avg'] or 1
    
    # График создания документов за последние 7 дней
    chart_data = []
    for i in range(6, -1, -1):
        date = today - timedelta(days=i)
        count = visible_docs.filter(created_at__date=date).count()
        chart_data.append({
            'date': date.strftime('%d.%m'),
            'count': count
        })
    chart_max = max((item['count'] for item in chart_data), default=1) or 1
    
    # Документы требующие внимания
    attention_needed = []
    
    # Просроченные документы пользователя
    overdue_docs = visible_docs.filter(
        Q(assigned_to=user) | Q(created_by=user),
        deadline__lt=today,
        status__in=['draft', 'sent_for_approval', 'coordination', 'approval', 'execution']
    )[:5]
    for doc in overdue_docs:
        attention_needed.append({
            'type': 'overdue',
            'document': doc,
            'message': f'Просрочен на {(today - doc.deadline).days} дн.'
        })
    
    # Документы на согласовании у пользователя
    if profile and profile.role in ['admin', 'manager']:
        from .models import WorkflowStep
        approval_docs = WorkflowStep.objects.filter(
            user=user,
            status='pending'
        ).select_related('document')[:5]
        for step in approval_docs:
            attention_needed.append({
                'type': 'approval',
                'document': step.document,
                'message': f'Ожидает вашего согласования (этап {step.step_number})'
            })
    
    context = {
        'total_documents': total_documents,
        'drafts': drafts,
        'sent_for_approval': sent_for_approval,
        'coordination': coordination,
        'approval': approval,
        'approved': approved,
        'execution': execution,
        'rejected': rejected,
        'returned': returned,
        'archived': archived,
        'in_approval_count': in_approval_count,
        'overdue': overdue,
        'urgent': urgent,
        'recent_documents': recent_documents,
        'documents_by_type': list(documents_by_type),
        'documents_by_status': list(documents_by_status),
        'documents_by_department': list(documents_by_department),
        'recent_activity': recent_activity,
        'pending_tasks': pending_tasks,
        'inprogress_tasks': inprogress_tasks,
        'overdue_tasks': overdue_tasks,
        'unread_notifications': unread_notifications,
        'signed_documents': signed_documents,
        'avg_version': round(avg_version, 1),
        'chart_data': chart_data,
        'chart_max': chart_max,
        'attention_needed': attention_needed,
        'can_see_all': can_see_all,
    }
    
    return render(request, 'documents/dashboard_analytics.html', context)


class DocumentListView(LoginRequiredMixin, ListView):
    """Список документов с фильтрацией"""
    model = Document
    template_name = 'documents/document_list_modern.html'
    context_object_name = 'documents'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = Document.objects.select_related('template', 'created_by', 'assigned_to')
        user = self.request.user
        
        # Права доступа в зависимости от роли
        if can_view_all_documents(user):
            # Админ, делопроизводитель и руководитель видят все документы
            pass
        else:
            # Сотрудник видит только:
            # - свои документы
            # - документы, назначенные ему
            queryset = queryset.filter(
                Q(created_by=user) | Q(assigned_to=user)
            )
        
        # Фильтрация
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(title__icontains=search) | 
                Q(registry_number__icontains=search) |
                Q(content__icontains=search)
            )
        
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        
        template_type = self.request.GET.get('template_type')
        if template_type:
            queryset = queryset.filter(template__type=template_type)
        
        date_from = self.request.GET.get('date_from')
        if date_from:
            queryset = queryset.filter(created_at__date__gte=date_from)
        
        date_to = self.request.GET.get('date_to')
        if date_to:
            queryset = queryset.filter(created_at__date__lte=date_to)
        
        only_overdue = self.request.GET.get('only_overdue')
        if only_overdue:
            queryset = queryset.filter(
                deadline__lt=timezone.now().date(),
                status__in=['draft', 'in_review']
            )
        
        return queryset.order_by('-created_at')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filter_form'] = DocumentFilterForm(self.request.GET)
        return context


class DocumentDetailView(LoginRequiredMixin, DetailView):
    """Детальная информация о документе"""
    model = Document
    template_name = 'documents/document_detail_modern.html'
    context_object_name = 'document'
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['history'] = self.object.history.select_related('user').all()
        context['workflow_steps'] = self.object.workflow_steps.select_related('user').all()
        context['approval_form'] = ApprovalForm()
        return context


class DocumentCreateView(LoginRequiredMixin, CreateView):
    """Создание нового документа"""
    model = Document
    form_class = DocumentForm
    template_name = 'documents/document_form_modern.html'
    success_url = reverse_lazy('documents:document_list')

    def dispatch(self, request, *args, **kwargs):
        """Проверка прав доступа на создание документов"""
        if not can_create_document(request.user):
            messages.error(request, 'У вас нет прав на создание документов. Обратитесь к делопроизводителю.')
            return redirect('documents:dashboard')
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs

    def get_initial(self):
        initial = super().get_initial()
        template_id = self.request.GET.get('template')
        if template_id:
            initial['template'] = template_id
        return initial

    def get_success_url(self):
        return reverse_lazy('documents:document_detail', kwargs={'pk': self.object.pk})

    def form_valid(self, form):
        form.instance.created_by = self.request.user
        response = super().form_valid(form)

        # Генерация файла из шаблона если шаблон выбран
        template = self.object.template
        if template and template.template_file:
            placeholder_values_json = self.request.POST.get('placeholder_values', '{}')
            try:
                user_placeholders = json.loads(placeholder_values_json)
            except (json.JSONDecodeError, ValueError):
                user_placeholders = {}

            from django.conf import settings as django_settings
            safe_title = re.sub(r'[<>:"/\\|?*]', '_', self.object.title)
            filename = f"{safe_title}_{self.object.id}.{template.file_format}"
            rel_output = os.path.join(
                'generated',
                str(timezone.now().year),
                str(timezone.now().month),
                filename
            )
            abs_output = os.path.join(str(django_settings.MEDIA_ROOT), rel_output)
            os.makedirs(os.path.dirname(abs_output), exist_ok=True)

            if template.file_format == 'pdf':
                # PDF: копируем шаблон как есть (замена плейсхолдеров в PDF без LibreOffice недоступна)
                import shutil
                shutil.copy2(template.template_file.path, abs_output)
                self.object.generated_file.name = rel_output
                self.object.save(update_fields=['generated_file'])
                messages.success(self.request, f'Документ "{self.object.title}" создан!')
            else:
                # docx / xlsx: заменяем плейсхолдеры
                replacements = {
                    'дата': timezone.now().strftime('%d.%m.%Y'),
                    'время': timezone.now().strftime('%H:%M'),
                    'название': self.object.title,
                }
                if template.placeholders and user_placeholders:
                    for ph in template.placeholders:
                        key = ph.get('name', '')
                        if key and key in user_placeholders:
                            val = user_placeholders[key]
                            replacements[key] = str(val) if val else ''

                success, error = generate_document_from_template(
                    template_file_path=template.template_file.path,
                    template_format=template.file_format,
                    output_path=abs_output,
                    replacements=replacements,
                )

                if success:
                    self.object.generated_file.name = rel_output
                    self.object.save(update_fields=['generated_file'])
                    messages.success(self.request, f'Документ "{self.object.title}" создан, файл сгенерирован!')
                else:
                    messages.warning(self.request, f'Документ создан, но ошибка генерации файла: {error}')
        else:
            messages.success(self.request, f'Документ "{self.object.title}" успешно создан!')

        # История и уведомления
        DocumentHistory.objects.create(
            document=self.object,
            user=self.request.user,
            action='Документ создан' + (f' из шаблона "{template.name}"' if template else '')
        )
        if self.object.assigned_to and self.object.assigned_to != self.request.user:
            send_document_assigned_email(self.object)

        return response


class DocumentUpdateView(LoginRequiredMixin, UpdateView):
    """Редактирование документа"""
    model = Document
    form_class = DocumentForm
    template_name = 'documents/document_form_modern.html'
    
    def dispatch(self, request, *args, **kwargs):
        """Проверка прав доступа на редактирование"""
        document = self.get_object()
        user_role = getattr(getattr(request.user, 'profile', None), 'role', None)
        if user_role in ['admin', 'clerk']:
            pass  # полный доступ
        else:
            # manager и employee могут редактировать только свои документы
            if document.created_by != request.user and document.assigned_to != request.user:
                messages.error(request, 'Вы можете редактировать только свои документы')
                return redirect('documents:document_detail', pk=document.pk)
        return super().dispatch(request, *args, **kwargs)
    
    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Добавляем информацию о том, что документ создан из шаблона
        if self.object.template:
            context['from_template'] = True
            context['template_name'] = self.object.template.name
        return context
    
    def get_success_url(self):
        return reverse_lazy('documents:document_detail', kwargs={'pk': self.object.pk})
    
    def form_valid(self, form):
        response = super().form_valid(form)
        
        # Создание записи в истории
        DocumentHistory.objects.create(
            document=self.object,
            user=self.request.user,
            action='Документ изменен'
        )
        
        messages.success(self.request, 'Документ успешно обновлен!')
        return response


class DocumentDeleteView(LoginRequiredMixin, ClerkRequiredMixin, DeleteView):
    """Удаление документа (только для делопроизводителя и администратора)"""
    model = Document
    template_name = 'documents/document_confirm_delete.html'
    success_url = reverse_lazy('documents:document_list')
    
    def dispatch(self, request, *args, **kwargs):
        """Проверка прав доступа на удаление"""
        document = self.get_object()
        if not can_delete_document(request.user, document):
            messages.error(request, 'У вас нет прав на удаление этого документа')
            return redirect('documents:document_detail', pk=document.pk)
        return super().dispatch(request, *args, **kwargs)
    
    def delete(self, request, *args, **kwargs):
        messages.success(request, 'Документ успешно удален!')
        return super().delete(request, *args, **kwargs)


@login_required
def create_from_template(request):
    """Создание документа из шаблона с генерацией файла — доступно всем"""
    if request.method == 'POST':
        # Получаем template из POST данных для создания формы с правильными полями
        template_id = request.POST.get('template')
        form = DocumentFromTemplateForm(request.POST, template_id=template_id)
        if form.is_valid():
            template = form.cleaned_data['template']
            
            # Подготовка замен для плейсхолдеров - теперь динамически
            replacements = {
                'дата': timezone.now().strftime('%d.%m.%Y'),
                'время': timezone.now().strftime('%H:%M'),
                'название': form.cleaned_data.get('title', ''),
            }
            
            # Добавляем кастомные плейсхолдеры из шаблона
            # placeholders is stored as a list of {name, label, type, ...} dicts
            raw_ph = template.placeholders or []
            if isinstance(raw_ph, dict):
                raw_ph = [{'name': k, **v} for k, v in raw_ph.items()]
            for item in raw_ph:
                key = item.get('name', '')
                if not key:
                    continue
                field_name = f'placeholder_{key}'
                value = form.cleaned_data.get(field_name, '')
                # Форматируем дату если это поле типа date
                if isinstance(value, date):
                    value = value.strftime('%d.%m.%Y')
                replacements[key] = str(value) if value else ''
            
            # Создание документа с защитой от дубликатов регистрационного номера
            max_save_attempts = 10
            document = None
            for save_attempt in range(max_save_attempts):
                try:
                    document = Document(
                        title=form.cleaned_data['title'],
                        template=template,
                        created_by=request.user,
                        assigned_to=form.cleaned_data.get('assigned_to'),
                        deadline=form.cleaned_data.get('deadline'),
                        status='draft'
                    )
                    
                    # Если статус не черновик, генерируем регистрационный номер заранее
                    if document.status != 'draft' and not document.registry_number:
                        document.registry_number = generate_unique_registry_number()
                    
                    # Пытаемся сохранить документ
                    document.save()
                    break  # Успешно сохранено
                    
                except IntegrityError as e:
                    if 'registry_number' in str(e) and save_attempt < max_save_attempts - 1:
                        # Если ошибка уникальности registry_number, пробуем снова
                        time.sleep(random.uniform(0.01, 0.05))
                        continue
                    else:
                        # Если другая ошибка или исчерпаны попытки, пробрасываем
                        raise
            
            if not document or not document.id:
                messages.error(request, 'Не удалось создать документ после нескольких попыток.')
                return redirect('documents:document-list')
            
            # Генерация файла из шаблона
            if template.template_file and template.file_format != 'html':
                # Путь к шаблону
                template_path = template.template_file.path
                
                # Генерация имени выходного файла (очистка от недопустимых символов)
                safe_title = re.sub(r'[<>:"/\\|?*]', '_', document.title)
                filename = f"{safe_title}_{document.id}.{template.file_format}"
                output_path = os.path.join('media', 'generated', str(timezone.now().year), 
                                          str(timezone.now().month), filename)
                
                # Создаем директорию если не существует
                os.makedirs(os.path.dirname(output_path), exist_ok=True)
                
                # Генерируем документ
                success, error = generate_document_from_template(
                    template_file_path=template_path,
                    template_format=template.file_format,
                    output_path=output_path,
                    replacements=replacements,
                )
                
                if success:
                    # Сохраняем путь к сгенерированному файлу
                    rel_path = os.path.relpath(output_path, 'media')
                    document.generated_file.name = rel_path
                    document.save()
                    messages.success(request, f'Документ "{document.title}" создан и файл сгенерирован!')
                else:
                    messages.warning(request, f'Документ создан, но ошибка генерации файла: {error}')
            
            elif template.html_template:
                # Старый способ - HTML шаблон
                content = template.html_template
                for key, value in replacements.items():
                    content = content.replace('{{' + key + '}}', value)
                document.content = content
                document.save()
                messages.success(request, f'Документ "{document.title}" создан из HTML шаблона!')
            
            else:
                messages.warning(request, 'Шаблон не содержит ни файла, ни HTML содержимого!')
            
            # История
            DocumentHistory.objects.create(
                document=document,
                user=request.user,
                action=f'Документ создан из шаблона "{template.name}"'
            )
            
            return redirect('documents:document_detail', pk=document.pk)
    else:
        # Получаем template_id из GET параметра если есть
        template_id = request.GET.get('template')
        form = DocumentFromTemplateForm(template_id=template_id)

    # Если шаблон передан через GET — блокируем выбор
    locked_template_id = None
    locked_template_name = None
    if request.method == 'GET' and request.GET.get('template'):
        try:
            _t = DocumentTemplate.objects.get(pk=request.GET['template'], is_active=True)
            locked_template_id = _t.pk
            locked_template_name = _t.name
        except DocumentTemplate.DoesNotExist:
            pass

    return render(request, 'documents/create_from_template.html', {
        'form': form,
        'locked_template_id': locked_template_id,
        'locked_template_name': locked_template_name,
    })


@login_required
def template_placeholders_json(request, pk):
    """AJAX: возвращает список плейсхолдеров шаблона в JSON"""
    from django.http import JsonResponse
    template = get_object_or_404(DocumentTemplate, pk=pk, is_active=True)
    phs = template.placeholders or []
    if isinstance(phs, dict):
        phs = [{'name': k, **v} for k, v in phs.items()]
    return JsonResponse({'placeholders': phs})


@login_required
def download_template_file(request, pk):
    """Скачивание файла шаблона"""
    template = get_object_or_404(DocumentTemplate, pk=pk)
    
    # Проверка наличия файла
    if not template.template_file:
        messages.error(request, 'У этого шаблона нет файла!')
        return redirect('documents:template_list')
    
    try:
        # Определяем имя файла
        filename = os.path.basename(template.template_file.name)
        
        # Открываем и отправляем файл
        response = FileResponse(template.template_file.open('rb'))
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    except Exception as e:
        messages.error(request, f'Ошибка при скачивании файла: {str(e)}')
        return redirect('documents:template_list')


@login_required
def generate_document_auto(request, template_id):
    """Автоматическая генерация документа из шаблона с заполнением данных"""
    from django.contrib.auth.models import User
    from datetime import datetime, timedelta
    import io
    from docx import Document as DocxDocument
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from django.core.files.base import ContentFile
    
    template = get_object_or_404(DocumentTemplate, pk=template_id)
    
    if request.method == 'GET':
        clerks = User.objects.filter(
            profile__role__in=['clerk', 'admin', 'manager']
        ).distinct()
        default_deadline = (datetime.now() + timedelta(days=7)).strftime('%Y-%m-%d')
        
        return render(request, 'documents/generate_document_auto.html', {
            'template': template,
            'clerks': clerks,
            'default_deadline': default_deadline
        })
    
    if request.method == 'POST':
        current_date = datetime.now().strftime('%d.%m.%Y')
        current_year = datetime.now().year
        
        # --- Собираем пользовательские плейсхолдеры ---
        user_placeholders = {}
        for key, value in request.POST.items():
            if key.startswith('placeholder_'):
                ph_name = key[len('placeholder_'):]
                user_placeholders[ph_name] = value
        
        # Формируем содержимое документа из плейсхолдеров
        if template.placeholders and user_placeholders:
            content_parts = []
            for ph in template.placeholders:
                val = user_placeholders.get(ph['name'], '')
                if val:
                    content_parts.append(f"{ph['label']}: {val}")
            content = '\n'.join(content_parts)
        else:
            content = request.POST.get('content', '')
        
        employee_name = request.user.get_full_name() or request.user.username
        
        # Автозаполнение данных в зависимости от типа документа
        replacements = {
            '{contract_number}': f"{random.randint(1, 999):03d}/{current_year}",
            '{date}': current_date,
            '{director_name}': "Шакирханов Рафаэль Ильдаревич",
            '{employee_name}': employee_name,
            '{position}': 'Менеджер',
            '{department}': 'Отдел продаж',
            '{start_date}': current_date,
            '{work_address}': "420030, Республика Татарстан, г Казань, тер. Снт Залив, д. 95",
            '{salary}': '50000',
            '{passport_data}': "4509 123456, выдан ОВД района Хамовники г. Москвы 01.01.2010",
            '{address}': "420030, Республика Татарстан, г Казань, тер. Снт Залив, д. 95",
            '{act_number}': f"{random.randint(1, 999):03d}/{current_year}",
            '{from_department}': "Отдел бухгалтерии",
            '{to_department}': "Отдел кадров",
            '{documents_list}': "1. Трудовой договор №001/2026\n2. Личная карточка Т-2\n3. Копия паспорта",
            '{total_count}': "5",
            '{total_pages}': "25",
            '{special_notes}': "Документы переданы в полном объеме, без повреждений.",
            '{from_person}': "Иванова И.И.",
            '{from_position}': "Главный бухгалтер",
            '{to_person}': "Сидорова С.С.",
            '{to_position}': "Начальник отдела кадров",
            '{order_number}': f"{random.randint(1, 999):03d}-к/{current_year}",
            '{contract_date}': current_date,
            '{probation_period}': "3",
            '{hr_person}': "Сидорова С.С.",
            '{department_head}': "Козлов К.К.",
            '{responsible_person}': "Сидорова С.С.",
            '{application_date}': current_date,
            '{memo_number}': f"{random.randint(1, 999):03d}",
            '{subject}': 'По вопросу согласования документов',
            '{content}': 'Довожу до Вашего сведения информацию о текущем состоянии дел.',
            '{requests}': "1. Рассмотреть данный вопрос\n2. Принять решение",
            '{additional_info}': "Прилагаю подробный отчет и расчеты.",
            '{attachments}': "1. Отчет о проделанной работе (5 стр.)",
            '{phone}': "+7 (495) 123-45-67",
            '{email}': request.user.email or "email@example.com",
            '{protocol_number}': f"{random.randint(1, 999):03d}",
            '{meeting_type}': "Собрание трудового коллектива",
            '{start_time}': "10:00",
            '{end_time}': "12:30",
            '{location}': "Конференц-зал, 3 этаж",
            '{chairman}': "Петров П.П.",
            '{secretary}': "Иванова И.И.",
            '{attendees}': "Петров П.П., Иванова И.И., Сидорова С.С., " + employee_name,
            '{agenda}': "1. Подведение итогов\n2. Утверждение плана\n3. Разное",
            '{speaker_1}': "Петров П.П.",
            '{speech_1}': "Доложил об итогах работы за прошедший квартал.",
            '{comments_1}': "Иванова И.И.: предложила усилить работу с клиентами",
            '{decision_1}': "Утвердить итоги работы за квартал.",
            '{speaker_2}': "Козлов К.К.",
            '{speech_2}': "Представил план работы на следующий квартал.",
            '{comments_2}': employee_name + ": предложил(а) внести дополнения в план",
            '{decision_2}': "Утвердить план работы с учетом внесенных предложений.",
            '{general_decisions}': "1. Утвердить итоги работы\n2. Утвердить план",
            '{tasks}': "1. Петрову П.П. - организовать встречу (срок: до 15.03.2026)",
            '{signatures}': "1. Иванова И.И. _____________\n2. Сидорова С.С. _____________",
            '{buyer_director}': "Шакирханов Рафаэль Ильдаревич",
            '{supplier_name}': 'ООО "Поставщик"',
            '{supplier_director}': "Иванов И.И.",
            '{supplier_basis}': "Устава",
            '{goods_list}': "1. Товар А - 100 шт. по цене 1000 руб. = 100 000 руб.",
            '{total_amount}': "300 000",
            '{total_amount_words}': "Триста тысяч",
            '{warranty_period}': "12",
            '{delivery_schedule}': "Партия 1: 100 000 руб. - до 15.03.2026",
            '{delivery_terms}': "доставка Поставщиком",
            '{delivery_address}': "420030, Республика Татарстан, г Казань, тер. Снт Залив, д. 95",
            '{payment_terms}': "Предоплата 30% в течение 5 банковских дней",
            '{end_date}': "31.12.2026",
            '{supplier_inn}': "7708123456",
            '{supplier_kpp}': "770801001",
            '{supplier_ogrn}': "1234567890987",
            '{supplier_address}': "г. Москва, ул. Садовая, д. 5",
            '{account_number}': "40702810100000012345",
            '{bank_name}': '"Сбербанк России" ПАО',
            '{correspondent_account}': "30101810400000000225",
            '{bik}': "044525225",
            '{supplier_account}': "40702810200000067890",
            '{supplier_bank}': '"ВТБ" ПАО',
            '{supplier_corr_account}': "30101810700000000187",
            '{supplier_bik}': "044525187",
            '{organization_name}': 'ООО "Первый ключ"',
            '{organization_inn}': "1683003976",
            '{organization_kpp}': "168301001",
            '{organization_ogrn}': "1221600022863",
            '{organization_address}': "420030, Республика Татарстан, г Казань, тер. Снт Залив, д. 95",
            '{organization_okpo}': "96615114",
            '{organization_okato}': "92401370000",
            '{organization_oktmo}': "92701000001",
        }
        
        # Добавляем пользовательские плейсхолдеры (перезаписывают автозаполнение)
        for name, value in user_placeholders.items():
            replacements['{' + name + '}'] = value
            replacements['{{' + name + '}}'] = value
        
        # Если есть файл шаблона — используем его, иначе берем html_template
        if template.template_file and template.file_format in ('docx',):
            try:
                from copy import deepcopy
                import re as _re
                
                template_file_path = template.template_file.path
                doc = DocxDocument(template_file_path)
                
                # Заменяем плейсхолдеры в параграфах
                def replace_in_para(para, replacements):
                    full_text = ''.join(run.text for run in para.runs)
                    new_text = full_text
                    for key, val in replacements.items():
                        new_text = new_text.replace(key, str(val))
                    if new_text != full_text:
                        for run in para.runs:
                            run.text = ''
                        if para.runs:
                            para.runs[0].text = new_text
                
                for para in doc.paragraphs:
                    replace_in_para(para, replacements)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_in_para(para, replacements)
                
                docx_io = io.BytesIO()
                doc.save(docx_io)
                docx_io.seek(0)
                file_ext = 'docx'
            except Exception as e:
                # Fallback: создать простой docx
                docx_io = None
                file_ext = 'docx'
        else:
            docx_io = None
            file_ext = 'docx'
        
        if not docx_io:
            # Создаём DOCX из текстового содержимого
            doc = DocxDocument()
            source_text = template.html_template or content
            for key, value in replacements.items():
                source_text = source_text.replace(key, str(value))
            
            paragraphs_list = source_text.split('\n') if source_text else [content]
            for i, para_text in enumerate(paragraphs_list):
                if para_text.strip():
                    paragraph = doc.add_paragraph()
                    stripped = para_text.strip()
                    # Заголовок: строка целиком написана капсом и короче 120 символов
                    is_header = (
                        stripped == stripped.upper()
                        and len(stripped) < 120
                        and any(c.isalpha() for c in stripped)
                    )
                    if is_header:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    run = paragraph.add_run(para_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(14)
                    run.font.color.rgb = None  # чёрный по умолчанию
                    paragraph.paragraph_format.line_spacing = 1.5
                    paragraph.paragraph_format.space_after = Pt(0)
            
            docx_io = io.BytesIO()
            doc.save(docx_io)
            docx_io.seek(0)
        
        # Генерируем уникальный регистрационный номер
        from django.db.models import Max
        last_doc = Document.objects.aggregate(Max('id'))
        next_id = (last_doc['id__max'] or 0) + 1
        registry_number = f"{next_id:05d}/{current_year}"
        
        # Получаем ответственного и срок исполнения из формы
        assigned_to_id = request.POST.get('assigned_to')
        assigned_to = User.objects.get(id=assigned_to_id) if assigned_to_id else None
        
        deadline_str = request.POST.get('deadline')
        deadline = datetime.strptime(deadline_str, '%Y-%m-%d').date() if deadline_str else None
        
        doc_title = request.POST.get('doc_title', f"{template.name} от {current_date}")
        
        # Создаем документ в БД
        document = Document.objects.create(
            title=doc_title,
            registry_number=registry_number,
            template=template,
            content=content,
            created_by=request.user,
            assigned_to=assigned_to,
            deadline=deadline,
            status='draft'
        )
        
        # Сохраняем файл
        filename = f"{template.name}_{document.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.generated_file.save(filename, ContentFile(docx_io.getvalue()), save=False)
        document.save()
        
        # История
        DocumentHistory.objects.create(
            document=document,
            user=request.user,
            action=f'Документ создан автоматически из шаблона "{template.name}"'
        )
        
        # Отправка email ответственному
        if assigned_to:
            send_document_assigned_email(document)
        
        messages.success(request, f'Документ "{document.title}" успешно создан!')
        return redirect('documents:document_detail', pk=document.id)
    
    # GET запрос - показываем форму
    users = User.objects.exclude(id=request.user.id)
    
    return render(request, 'documents/generate_document_auto.html', {
        'template': template,
        'users': users
    })


@login_required
def download_generated_file(request, pk):
    """Скачивание сгенерированного файла документа"""
    document = get_object_or_404(Document, pk=pk)
    
    # Проверка прав доступа
    if not (document.created_by == request.user or 
            document.assigned_to == request.user or
            (hasattr(request.user, 'profile') and request.user.profile.is_clerk)):
        messages.error(request, 'У вас нет прав для скачивания этого документа!')
        return redirect('documents:document_detail', pk=pk)
    
    # Проверка наличия файла
    if not document.generated_file:
        messages.error(request, 'У этого документа нет сгенерированного файла!')
        return redirect('documents:document_detail', pk=pk)
    
    try:
        # Открываем файл для скачивания
        file_path = document.generated_file.path
        
        if not os.path.exists(file_path):
            messages.error(request, 'Файл не найден на сервере!')
            return redirect('documents:document_detail', pk=pk)
        
        # Определяем MIME-тип по расширению
        ext = os.path.splitext(file_path)[1].lower()
        content_types = {
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            '.pdf': 'application/pdf',
        }
        content_type = content_types.get(ext, 'application/octet-stream')
        
        # Возвращаем файл
        response = FileResponse(open(file_path, 'rb'), content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{os.path.basename(file_path)}"'
        
        # Логируем скачивание
        DocumentHistory.objects.create(
            document=document,
            user=request.user,
            action='Файл документа скачан'
        )
        
        return response
    
    except Exception as e:
        messages.error(request, f'Ошибка при скачивании файла: {str(e)}')
        return redirect('documents:document_detail', pk=pk)


@login_required
def register_document(request, pk):
    """Регистрация документа (присвоение номера) — доступна всем пользователям"""
    document = get_object_or_404(Document, pk=pk)
    
    if document.status == 'draft':
        document.status = 'in_review'
        if not document.registry_number:
            document.registry_number = generate_unique_registry_number()
        document.save()
        
        DocumentHistory.objects.create(
            document=document,
            user=request.user,
            action=f'Документ зарегистрирован под номером {document.registry_number}'
        )
        
        # Уведомление ответственному
        if document.assigned_to:
            Notification.objects.create(
                user=document.assigned_to,
                message=f'Вам назначен документ "{document.title}" ({document.registry_number})',
                document=document
            )
            # Отправка email
            send_document_assigned_email(document)
        
        messages.success(request, f'Документ зарегистрирован под номером {document.registry_number}')
    else:
        messages.warning(request, 'Документ уже зарегистрирован!')
    
    return redirect('documents:document_detail', pk=pk)


@login_required
@clerk_required
def setup_workflow(request, pk):
    """Настройка маршрута согласования (делопроизводитель и администратор)"""
    document = get_object_or_404(Document, pk=pk)
    
    if request.method == 'POST':
        form = WorkflowRouteForm(request.POST)
        if form.is_valid():
            approvers = form.cleaned_data['approvers']
            
            # Удаляем старые этапы
            document.workflow_steps.all().delete()
            
            # Создаем новые этапы
            for i, approver in enumerate(approvers, start=1):
                WorkflowStep.objects.create(
                    document=document,
                    step_number=i,
                    user=approver,
                    status='pending'
                )
            
            document.current_step = 0
            document.save()
            
            DocumentHistory.objects.create(
                document=document,
                user=request.user,
                action=f'Настроен маршрут согласования: {len(approvers)} этапов'
            )
            
            # Уведомляем первого согласующего
            if approvers:
                first_step = document.workflow_steps.filter(step_number=1).first()
                if first_step:
                    Notification.objects.create(
                        user=first_step.user,
                        message=f"Документ '{document.title}' ожидает вашего согласования (этап 1)",
                        document=document
                    )
                    # Отправка email
                    send_workflow_step_notification(first_step, document)
            
            messages.success(request, 'Маршрут согласования настроен!')
            return redirect('documents:document_detail', pk=pk)
    else:
        form = WorkflowRouteForm()
    
    return render(request, 'documents/setup_workflow.html', {'form': form, 'document': document})


@login_required
def approve_workflow_step(request, step_id):
    """Согласование этапа маршрута документа"""
    from django.utils import timezone
    
    step = get_object_or_404(WorkflowStep, pk=step_id)
    document = step.document
    
    # admin и clerk могут согласовать любой этап; остальные — только назначенные
    user_role = getattr(getattr(request.user, 'profile', None), 'role', None)
    if step.user != request.user and user_role not in ['admin', 'clerk']:
        messages.error(request, 'У вас нет прав на согласование этого этапа!')
        return redirect('documents:document_detail', pk=document.id)
    
    # Проверка что этап еще не завершен
    if step.status != 'pending':
        messages.warning(request, 'Этот этап уже обработан!')
        return redirect('documents:document_detail', pk=document.id)
    
    # Проверка что это текущий активный этап (предыдущие должны быть согласованы)
    previous_steps = document.workflow_steps.filter(step_number__lt=step.step_number)
    if previous_steps.filter(status='pending').exists():
        messages.error(request, 'Необходимо дождаться согласования предыдущих этапов!')
        return redirect('documents:document_detail', pk=document.id)
    
    if previous_steps.filter(status='rejected').exists():
        messages.error(request, 'Документ был отклонен на предыдущем этапе!')
        return redirect('documents:document_detail', pk=document.id)
    
    if request.method == 'POST':
        action = request.POST.get('action')
        comment = request.POST.get('comment', '')
        
        # Обновляем статус этапа
        if action == 'approve':
            step.status = 'approved'
            step.comment = comment
            step.completed_at = timezone.now()
            step.save()
            
            # Проверяем, есть ли еще этапы
            next_steps = document.workflow_steps.filter(
                step_number__gt=step.step_number,
                status='pending'
            ).order_by('step_number')
            
            if next_steps.exists():
                # Уведомляем следующего согласующего
                next_step = next_steps.first()
                Notification.objects.create(
                    user=next_step.user,
                    message=f"Документ '{document.title}' ожидает вашего согласования (этап {next_step.step_number})"
                )
                # Отправка email
                send_workflow_step_notification(next_step, document)
                messages.success(request, f'Этап согласован! Документ передан на этап {next_step.step_number}')
            else:
                # Все этапы пройдены - документ согласован
                document.status = 'approved'
                document.save()
                Notification.objects.create(
                    user=document.created_by,
                    message=f"Документ '{document.title}' полностью согласован!"
                )
                # Отправка email
                send_document_approved_email(document, request.user)
                messages.success(request, 'Этап согласован! Документ полностью согласован!')
            
            # История
            DocumentHistory.objects.create(
                document=document,
                user=request.user,
                action=f"Согласовал этап {step.step_number}",
                comment=comment
            )
            
        elif action == 'reject':
            step.status = 'rejected'
            step.comment = comment
            step.completed_at = timezone.now()
            step.save()
            
            # Отклоняем документ
            document.status = 'rejected'
            document.save()
            
            # Уведомляем создателя
            Notification.objects.create(
                user=document.created_by,
                message=f"Документ '{document.title}' отклонен на этапе {step.step_number} пользователем {request.user.get_full_name()}"
            )
            # Отправка email
            send_document_rejected_email(document, request.user, comment)
            
            # История
            DocumentHistory.objects.create(
                document=document,
                user=request.user,
                action=f"Отклонил документ на этапе {step.step_number}",
                comment=comment
            )
            
            messages.warning(request, 'Документ отклонен!')
    
    return redirect('documents:document_detail', pk=document.id)


@login_required
@clerk_or_manager_required
def approve_document(request, pk):
    """Утверждение/отклонение документа — для делопроизводителя, руководителя и администратора"""
    document = get_object_or_404(Document, pk=pk)
    
    if request.method == 'POST':
        form = ApprovalForm(request.POST)
        if form.is_valid():
            action = form.cleaned_data['action']
            comment = form.cleaned_data['comment']
            
            if action == 'approve':
                document.status = 'approved'
                action_text = 'Документ утвержден'
                messages.success(request, 'Документ успешно утвержден!')
            else:
                document.status = 'rejected'
                action_text = 'Документ отклонен'
                messages.warning(request, 'Документ отклонен!')
            
            document.save()
            
            # История
            DocumentHistory.objects.create(
                document=document,
                user=request.user,
                action=action_text,
                comment=comment
            )
            
            # Уведомление создателю
            Notification.objects.create(
                user=document.created_by,
                message=f'Документ "{document.title}" ({document.registry_number}) {action_text.lower()}',
                document=document
            )
            # Отправка email
            if action == 'approve':
                send_document_approved_email(document, request.user)
            elif action == 'reject':
                send_document_rejected_email(document, request.user, comment)
            
            return redirect('documents:document_detail', pk=pk)
    
    return redirect('documents:document_detail', pk=pk)


@login_required
@clerk_required
def archive_document(request, pk):
    """Архивирование документа (делопроизводитель и администратор)"""
    document = get_object_or_404(Document, pk=pk)
    
    document.status = 'archived'
    document.save()
    
    DocumentHistory.objects.create(
        document=document,
        user=request.user,
        action='Документ перемещен в архив'
    )
    
    messages.success(request, 'Документ перемещен в архив!')
    return redirect('documents:document_list')


@login_required
@clerk_required
def bulk_archive(request):
    """Массовое архивирование документов (делопроизводитель и администратор)"""
    if request.method == 'POST':
        document_ids = request.POST.getlist('document_ids')
        
        documents = Document.objects.filter(id__in=document_ids)
        count = documents.update(status='archived')
        
        for doc in documents:
            DocumentHistory.objects.create(
                document=doc,
                user=request.user,
                action='Документ перемещен в архив (массовая операция)'
            )
        
        return JsonResponse({'success': True, 'count': count})
    
    return JsonResponse({'success': False})


# Template Views

class TemplateListView(LoginRequiredMixin, ListView):
    """Список шаблонов — доступен всем авторизованным"""
    model = DocumentTemplate
    template_name = 'documents/template_list_modern.html'
    context_object_name = 'templates'
    paginate_by = 20
    
    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Фильтрация по поиску
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(
                Q(name__icontains=search) | Q(description__icontains=search)
            )
        
        # Фильтрация по типу
        template_type = self.request.GET.get('type')
        if template_type:
            queryset = queryset.filter(type=template_type)
        
        # Фильтрация по формату
        file_format = self.request.GET.get('format')
        if file_format:
            queryset = queryset.filter(file_format=file_format)
        
        return queryset.order_by('-created_at')
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Статистика
        all_templates = DocumentTemplate.objects.all()
        context['types_count'] = all_templates.values('type').distinct().count()
        context['formats_count'] = all_templates.values('file_format').distinct().count()
        
        return context


class TemplateCreateView(LoginRequiredMixin, CreateView):
    """Создание шаблона — доступно всем авторизованным"""
    model = DocumentTemplate
    form_class = DocumentTemplateForm
    template_name = 'documents/template_form_modern.html'
    success_url = reverse_lazy('documents:template_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['placeholders_json_str'] = '[]'
        return context

    def form_valid(self, form):
        placeholders_json = self.request.POST.get('placeholders_json', '[]')
        try:
            form.instance.placeholders = json.loads(placeholders_json)
        except (json.JSONDecodeError, ValueError):
            form.instance.placeholders = []
        messages.success(self.request, 'Шаблон успешно создан!')
        return super().form_valid(form)


class TemplateUpdateView(LoginRequiredMixin, UpdateView):
    """Редактирование шаблона — доступно всем авторизованным"""
    model = DocumentTemplate
    form_class = DocumentTemplateForm
    template_name = 'documents/template_form_modern.html'
    success_url = reverse_lazy('documents:template_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['placeholders_json_str'] = json.dumps(
            self.object.placeholders if self.object and self.object.placeholders else []
        )
        return context

    def form_valid(self, form):
        placeholders_json = self.request.POST.get('placeholders_json', '[]')
        try:
            form.instance.placeholders = json.loads(placeholders_json)
        except (json.JSONDecodeError, ValueError):
            form.instance.placeholders = []
        messages.success(self.request, 'Шаблон успешно обновлен!')
        return super().form_valid(form)


class TemplateDeleteView(LoginRequiredMixin, ClerkRequiredMixin, DeleteView):
    """Удаление шаблона"""
    model = DocumentTemplate
    template_name = 'documents/template_confirm_delete.html'
    success_url = reverse_lazy('documents:template_list')


@login_required
@manager_required
def reports(request):
    """Отчеты и аналитика (руководитель и администратор)"""
    from django.contrib.auth.models import User as AuthUser

    today_date = timezone.now().date()

    # ── Period selection ──────────────────────────────────────
    preset = request.GET.get('preset', '30')
    date_from_str = request.GET.get('date_from', '')
    date_to_str = request.GET.get('date_to', '')

    PRESETS = {
        '7':   ('7 дней',    today_date - timedelta(days=6),   today_date),
        '30':  ('30 дней',   today_date - timedelta(days=29),  today_date),
        '90':  ('90 дней',   today_date - timedelta(days=89),  today_date),
        '365': ('Год',       today_date - timedelta(days=364), today_date),
    }

    if date_from_str or date_to_str:
        preset = 'custom'
        try:
            date_from = datetime.strptime(date_from_str, '%Y-%m-%d').date() if date_from_str else None
            date_to   = datetime.strptime(date_to_str,   '%Y-%m-%d').date() if date_to_str else today_date
        except ValueError:
            date_from, date_to = today_date - timedelta(days=29), today_date
        period_label = f"{date_from.strftime('%d.%m.%Y') if date_from else '...'} – {date_to.strftime('%d.%m.%Y')}"
    elif preset in PRESETS:
        period_label, date_from, date_to = PRESETS[preset]
    else:
        preset = '30'
        period_label, date_from, date_to = PRESETS['30']

    def period_filter(qs, field='created_at'):
        if date_from:
            qs = qs.filter(**{f'{field}__date__gte': date_from})
        if date_to:
            qs = qs.filter(**{f'{field}__date__lte': date_to})
        return qs

    all_docs = Document.objects.all()
    period_docs = period_filter(all_docs)

    # ── KPI summary ───────────────────────────────────────────
    total_all    = all_docs.count()
    total_period = period_docs.count()
    approved_period  = period_filter(all_docs.filter(status='approved')).count()
    rejected_period  = period_filter(all_docs.filter(status='rejected')).count()
    in_work_period   = period_docs.exclude(status__in=['draft', 'archived', 'rejected']).count()
    in_review_period = period_docs.filter(status__in=['in_review', 'sent_for_approval', 'coordination', 'approval']).count()
    overdue_all = all_docs.filter(
        deadline__lt=today_date,
        status__in=['draft', 'sent_for_approval', 'coordination', 'approval', 'execution']
    )

    # ── By status ─────────────────────────────────────────────
    by_status = period_docs.values('status').annotate(count=Count('id')).order_by('-count')

    # ── By type ───────────────────────────────────────────────
    by_type = period_docs.filter(template__isnull=False).values(
        'template__type'
    ).annotate(count=Count('id')).order_by('-count')
    no_template_count = period_docs.filter(template__isnull=True).count()

    # ── By user (top creators) ────────────────────────────────
    by_creator = period_docs.values(
        'created_by__id', 'created_by__first_name', 'created_by__last_name', 'created_by__username'
    ).annotate(count=Count('id')).order_by('-count')[:10]

    # ── Daily dynamics ────────────────────────────────────────
    if date_from and date_to:
        delta_days = (date_to - date_from).days + 1
    else:
        delta_days = 30
    # cap at 90 points for readability
    if delta_days > 90:
        # group by week instead
        daily_data = []
        group_by = 'week'
        from django.db.models.functions import TruncWeek
        weekly = period_docs.annotate(wk=TruncWeek('created_at')).values('wk').annotate(count=Count('id')).order_by('wk')
        for row in weekly:
            daily_data.append({'label': row['wk'].strftime('%d.%m'), 'count': row['count']})
        period_unit = 'по неделям'
    else:
        group_by = 'day'
        period_unit = 'по дням'
        daily_counts = {
            row['day']: row['count']
            for row in period_docs.annotate(day=TruncDate('created_at')).values('day').annotate(count=Count('id'))
        } if date_from else {}
        daily_data = []
        if date_from and date_to:
            cur = date_from
            while cur <= date_to:
                daily_data.append({'label': cur.strftime('%d.%m'), 'count': daily_counts.get(cur, 0)})
                cur += timedelta(days=1)

    # ── Workflow approval stats ───────────────────────────────
    try:
        from .models import WorkflowApproval
        wa_qs = WorkflowApproval.objects.all()
        if date_from:
            wa_qs = wa_qs.filter(decision_date__date__gte=date_from)
        if date_to:
            wa_qs = wa_qs.filter(decision_date__date__lte=date_to)
        wf_approved = wa_qs.filter(decision='approved').count()
        wf_rejected = wa_qs.filter(decision='rejected').count()
        wf_pending  = wa_qs.filter(decision='pending').count()
        wf_total    = wf_approved + wf_rejected + wf_pending
    except Exception:
        wf_approved = wf_rejected = wf_pending = wf_total = 0

    # ── Avg processing time (created_at → updated_at for approved docs) ──
    try:
        avg_qs = period_filter(all_docs.filter(status='approved'), field='created_at')
        avg_duration = avg_qs.annotate(
            duration=ExpressionWrapper(F('updated_at') - F('created_at'), output_field=DurationField())
        ).aggregate(avg=Avg('duration'))['avg']
        avg_days = round(avg_duration.days + avg_duration.seconds / 86400, 1) if avg_duration else None
    except Exception:
        avg_days = None

    # ── Top assignees ─────────────────────────────────────────
    by_assignee = period_docs.filter(assigned_to__isnull=False).values(
        'assigned_to__first_name', 'assigned_to__last_name', 'assigned_to__username'
    ).annotate(count=Count('id')).order_by('-count')[:10]

    # ── Document type display map ─────────────────────────────
    TYPE_LABELS = dict(DocumentTemplate.TYPE_CHOICES)

    context = {
        # period
        'preset': preset,
        'presets': PRESETS,
        'date_from': date_from.strftime('%Y-%m-%d') if date_from else '',
        'date_to':   date_to.strftime('%Y-%m-%d')   if date_to   else '',
        'period_label': period_label,
        'period_unit': period_unit,
        'today': timezone.now(),
        # kpi
        'total_all': total_all,
        'total_period': total_period,
        'approved_period': approved_period,
        'rejected_period': rejected_period,
        'in_work_period': in_work_period,
        'in_review_period': in_review_period,
        'overdue_count': overdue_all.count(),
        'avg_days': avg_days,
        # tables / charts
        'by_status': list(by_status),
        'by_type': list(by_type),
        'no_template_count': no_template_count,
        'by_creator': list(by_creator),
        'by_assignee': list(by_assignee),
        'daily_data': daily_data,
        # workflow
        'wf_approved': wf_approved,
        'wf_rejected': wf_rejected,
        'wf_pending':  wf_pending,
        'wf_total':    wf_total,
        # overdue table
        'overdue_documents': overdue_all.select_related('assigned_to', 'created_by').order_by('deadline')[:30],
        'type_labels': TYPE_LABELS,
    }

    # ── Narrative text for print ──────────────────────────────
    STATUS_LABELS_PRINT = dict(Document.STATUS_CHOICES)
    _narrative_d = {
        'period_label': period_label,
        'today': timezone.now(),
        'total_period': total_period,
        'total_all': total_all,
        'approved_period': approved_period,
        'rejected_period': rejected_period,
        'overdue_count': overdue_all.count(),
        'in_work_period': in_work_period,
        'in_review_period': in_review_period,
        'avg_days': avg_days,
        'by_status': list(by_status),
        'STATUS_LABELS': STATUS_LABELS_PRINT,
        'by_type': list(by_type),
        'TYPE_LABELS': TYPE_LABELS,
        'wf_total': wf_total,
        'wf_approved': wf_approved,
        'wf_rejected': wf_rejected,
        'wf_pending': wf_pending,
        'by_creator': list(by_creator),
        'by_assignee': list(by_assignee),
        'overdue_list': list(overdue_all.select_related('assigned_to', 'created_by').order_by('deadline')[:30]),
    }
    context['narrative_sections'] = _report_build_narrative(_narrative_d)

    return render(request, 'documents/reports.html', context)


# ─────────────────────────────────────────────────────────────────────────────
# Helper: shared data builder used by all export views
# ─────────────────────────────────────────────────────────────────────────────
def _report_get_data(request):
    """Собирает данные отчёта для экспорта (Excel / PDF / DOCX)."""
    today_date = timezone.now().date()

    preset = request.GET.get('preset', '30')
    date_from_str = request.GET.get('date_from', '')
    date_to_str   = request.GET.get('date_to', '')

    PRESETS = {
        '7':   ('7 дней',    today_date - timedelta(days=6),   today_date),
        '30':  ('30 дней',   today_date - timedelta(days=29),  today_date),
        '90':  ('90 дней',   today_date - timedelta(days=89),  today_date),
        '365': ('Год',       today_date - timedelta(days=364), today_date),
    }

    if date_from_str or date_to_str:
        try:
            date_from = datetime.strptime(date_from_str, '%Y-%m-%d').date() if date_from_str else None
            date_to   = datetime.strptime(date_to_str,   '%Y-%m-%d').date() if date_to_str else today_date
        except ValueError:
            date_from, date_to = today_date - timedelta(days=29), today_date
        period_label = f"{date_from.strftime('%d.%m.%Y') if date_from else '...'} – {date_to.strftime('%d.%m.%Y')}"
    elif preset in PRESETS:
        period_label, date_from, date_to = PRESETS[preset]
    else:
        preset = '30'
        period_label, date_from, date_to = PRESETS['30']

    def pf(qs, field='created_at'):
        if date_from:
            qs = qs.filter(**{f'{field}__date__gte': date_from})
        if date_to:
            qs = qs.filter(**{f'{field}__date__lte': date_to})
        return qs

    all_docs    = Document.objects.all()
    period_docs = pf(all_docs)

    overdue_all = all_docs.filter(
        deadline__lt=today_date,
        status__in=['draft', 'sent_for_approval', 'coordination', 'approval', 'execution']
    )

    try:
        avg_qs = pf(all_docs.filter(status='approved'))
        avg_dur = avg_qs.annotate(
            dur=ExpressionWrapper(F('updated_at') - F('created_at'), output_field=DurationField())
        ).aggregate(avg=Avg('dur'))['avg']
        avg_days = round(avg_dur.days + avg_dur.seconds / 86400, 1) if avg_dur else None
    except Exception:
        avg_days = None

    try:
        from .models import WorkflowApproval
        wa_qs = WorkflowApproval.objects.all()
        if date_from:
            wa_qs = wa_qs.filter(decision_date__date__gte=date_from)
        if date_to:
            wa_qs = wa_qs.filter(decision_date__date__lte=date_to)
        wf_approved = wa_qs.filter(decision='approved').count()
        wf_rejected = wa_qs.filter(decision='rejected').count()
        wf_pending  = wa_qs.filter(decision='pending').count()
        wf_total    = wf_approved + wf_rejected + wf_pending
    except Exception:
        wf_approved = wf_rejected = wf_pending = wf_total = 0

    STATUS_LABELS = dict(Document.STATUS_CHOICES)
    TYPE_LABELS   = dict(DocumentTemplate.TYPE_CHOICES)

    by_status = list(period_docs.values('status').annotate(count=Count('id')).order_by('-count'))
    by_type   = list(period_docs.filter(template__isnull=False).values('template__type').annotate(count=Count('id')).order_by('-count'))
    by_creator = list(period_docs.values(
        'created_by__first_name', 'created_by__last_name', 'created_by__username'
    ).annotate(count=Count('id')).order_by('-count')[:10])
    by_assignee = list(period_docs.filter(assigned_to__isnull=False).values(
        'assigned_to__first_name', 'assigned_to__last_name', 'assigned_to__username'
    ).annotate(count=Count('id')).order_by('-count')[:10])
    overdue_list = list(overdue_all.select_related('assigned_to', 'created_by').order_by('deadline')[:50])

    return {
        'period_label': period_label,
        'date_from': date_from, 'date_to': date_to,
        'today': timezone.now(),
        'total_all': all_docs.count(),
        'total_period': period_docs.count(),
        'approved_period': pf(all_docs.filter(status='approved')).count(),
        'rejected_period': pf(all_docs.filter(status='rejected')).count(),
        'in_work_period': period_docs.exclude(status__in=['draft', 'archived', 'rejected']).count(),
        'in_review_period': period_docs.filter(status__in=['in_review', 'sent_for_approval', 'coordination', 'approval']).count(),
        'overdue_count': overdue_all.count(),
        'avg_days': avg_days,
        'by_status': by_status,
        'by_type': by_type,
        'by_creator': by_creator,
        'by_assignee': by_assignee,
        'overdue_list': overdue_list,
        'wf_approved': wf_approved, 'wf_rejected': wf_rejected,
        'wf_pending': wf_pending, 'wf_total': wf_total,
        'STATUS_LABELS': STATUS_LABELS, 'TYPE_LABELS': TYPE_LABELS,
    }


# ─────────────────────────────────────────────────────────────────────────────
# Report export: Excel
# ─────────────────────────────────────────────────────────────────────────────
@login_required
@manager_required
def report_export_excel(request):
    """Экспорт аналитического отчёта в Excel (.xlsx)"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter
    from io import BytesIO

    d = _report_get_data(request)

    wb = openpyxl.Workbook()

    HDR_FILL  = PatternFill('solid', fgColor='3730A3')   # indigo-800
    HDR_FONT  = Font(color='FFFFFF', bold=True, size=11)
    SUB_FILL  = PatternFill('solid', fgColor='E0E7FF')   # indigo-100
    SUB_FONT  = Font(color='312E81', bold=True, size=10)
    THIN      = Side(style='thin', color='D1D5DB')
    BORDER    = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
    CENTER    = Alignment(horizontal='center', vertical='center', wrap_text=True)
    LEFT      = Alignment(horizontal='left', vertical='center')

    def set_hdr(ws, row, cols, title):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=cols)
        c = ws.cell(row=row, column=1, value=title)
        c.fill = HDR_FILL; c.font = HDR_FONT; c.alignment = CENTER
        ws.row_dimensions[row].height = 22

    def set_col_hdr(ws, row, values):
        for col, v in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=v)
            c.fill = SUB_FILL; c.font = SUB_FONT; c.alignment = CENTER
            c.border = BORDER

    def write_row(ws, row, values, bold=False):
        for col, v in enumerate(values, 1):
            c = ws.cell(row=row, column=col, value=v)
            c.border = BORDER
            c.alignment = LEFT
            if bold:
                c.font = Font(bold=True)

    def autofit(ws):
        for col_cells in ws.columns:
            length = max((len(str(c.value or '')) for c in col_cells), default=8)
            ws.column_dimensions[get_column_letter(col_cells[0].column)].width = min(length + 4, 50)

    # ── Sheet 1: Сводка ──────────────────────────────────────
    ws1 = wb.active
    ws1.title = 'Сводка'
    set_hdr(ws1, 1, 2, f'Аналитический отчёт — {d["period_label"]}')
    set_hdr(ws1, 2, 2, f'Сформирован: {d["today"].strftime("%d.%m.%Y %H:%M")}')
    ws1.row_dimensions[2].height = 16
    set_col_hdr(ws1, 3, ['Показатель', 'Значение'])
    rows = [
        ('Всего документов в системе',         d['total_all']),
        ('Документов за период',               d['total_period']),
        ('Утверждено за период',               d['approved_period']),
        ('Отклонено за период',                d['rejected_period']),
        ('В работе за период',                  d['in_work_period']),
        ('На рассмотрении за период',           d['in_review_period']),
        ('Просрочено (активных)',              d['overdue_count']),
        ('Среднее время обработки (дней)',     d['avg_days'] if d['avg_days'] is not None else '—'),
        ('Согласований: утверждено',           d['wf_approved']),
        ('Согласований: отклонено',            d['wf_rejected']),
        ('Согласований: ожидает',              d['wf_pending']),
    ]
    for i, (k, v) in enumerate(rows, 4):
        write_row(ws1, i, [k, v])
    autofit(ws1)

    # ── Sheet 2: По статусам ─────────────────────────────────
    ws2 = wb.create_sheet('По статусам')
    set_hdr(ws2, 1, 3, f'Разбивка по статусам — {d["period_label"]}')
    set_col_hdr(ws2, 2, ['Статус', 'Количество', 'Доля, %'])
    total_p = d['total_period'] or 1
    for i, item in enumerate(d['by_status'], 3):
        label = d['STATUS_LABELS'].get(item['status'], item['status'])
        pct   = round(item['count'] / total_p * 100, 1)
        write_row(ws2, i, [label, item['count'], pct])
    autofit(ws2)

    # ── Sheet 3: По типам ────────────────────────────────────
    ws3 = wb.create_sheet('По типам')
    set_hdr(ws3, 1, 3, f'Разбивка по типам — {d["period_label"]}')
    set_col_hdr(ws3, 2, ['Тип документа', 'Количество', 'Доля, %'])
    for i, item in enumerate(d['by_type'], 3):
        label = d['TYPE_LABELS'].get(item['template__type'], item['template__type'] or '—')
        pct   = round(item['count'] / total_p * 100, 1)
        write_row(ws3, i, [label, item['count'], pct])
    autofit(ws3)

    # ── Sheet 4: Авторы ──────────────────────────────────────
    ws4 = wb.create_sheet('Топ авторов')
    set_hdr(ws4, 1, 3, f'Топ авторов документов — {d["period_label"]}')
    set_col_hdr(ws4, 2, ['#', 'Сотрудник', 'Создано документов'])
    for i, item in enumerate(d['by_creator'], 3):
        fn = item.get('created_by__first_name', '')
        ln = item.get('created_by__last_name', '')
        un = item.get('created_by__username', '')
        name = f'{fn} {ln}'.strip() or un
        write_row(ws4, i, [i - 2, name, item['count']])
    autofit(ws4)

    # ── Sheet 5: Исполнители ─────────────────────────────────
    ws5 = wb.create_sheet('Топ исполнителей')
    set_hdr(ws5, 1, 3, f'Топ исполнителей — {d["period_label"]}')
    set_col_hdr(ws5, 2, ['#', 'Сотрудник', 'Назначено документов'])
    for i, item in enumerate(d['by_assignee'], 3):
        fn = item.get('assigned_to__first_name', '')
        ln = item.get('assigned_to__last_name', '')
        un = item.get('assigned_to__username', '')
        name = f'{fn} {ln}'.strip() or un
        write_row(ws5, i, [i - 2, name, item['count']])
    autofit(ws5)

    # ── Sheet 6: Просроченные ────────────────────────────────
    ws6 = wb.create_sheet('Просроченные')
    set_hdr(ws6, 1, 5, 'Просроченные документы (активные)')
    set_col_hdr(ws6, 2, ['Рег. номер', 'Название', 'Срок', 'Ответственный', 'Статус'])
    for i, doc in enumerate(d['overdue_list'], 3):
        assignee = doc.assigned_to.get_full_name() if doc.assigned_to else '—'
        deadline = doc.deadline.strftime('%d.%m.%Y') if doc.deadline else '—'
        status   = d['STATUS_LABELS'].get(doc.status, doc.status)
        write_row(ws6, i, [doc.registry_number or '—', doc.title, deadline, assignee, status])
    autofit(ws6)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    fname = f'report_{d["today"].strftime("%Y%m%d")}.xlsx'
    resp = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    return resp


def _report_build_narrative(d):
    """
    Формирует структурированный аналитический текст по данным отчёта.
    Возвращает список словарей: {'title': str, 'body': str}
    """
    sections = []
    total_p  = d['total_period']
    total_a  = d['total_all']

    # ── 1. Вступление / резюме ────────────────────────────────────────────────
    intro_body = (
        f"Настоящий аналитический отчёт охватывает период «{d['period_label']}» и составлен "
        f"{d['today'].strftime('%d.%m.%Y')} в {d['today'].strftime('%H:%M')}. "
        f"В указанный период система электронного документооборота зафиксировала "
        f"{total_p} {'документ' if total_p == 1 else 'документа' if 2 <= total_p <= 4 else 'документов'} "
        f"(всего в системе накоплено {total_a} "
        f"{'документ' if total_a == 1 else 'документа' if 2 <= total_a <= 4 else 'документов'}). "
    )
    if total_p == 0:
        intro_body += (
            "За анализируемый период документооборот отсутствовал, что может свидетельствовать "
            "о нерабочем периоде, праздничных днях или временной остановке бизнес-процессов."
        )
    elif total_p < 5:
        intro_body += (
            "Документопоток за период отличается низкой интенсивностью. "
            "Рекомендуется сопоставить показатели с аналогичным периодом прошлого года для оценки тенденции."
        )
    else:
        intro_body += (
            "Документопоток за период находится в активной фазе. "
            "В ходе дальнейшего анализа рассматриваются ключевые показатели эффективности, "
            "структура документооборота по статусам и типам, результаты процессов согласования, "
            "а также нагрузка на участников процесса."
        )
    sections.append({'title': 'Введение', 'body': intro_body})

    # ── 2. Анализ КПЭ ────────────────────────────────────────────────────────
    if total_p > 0:
        approval_pct = round(d['approved_period'] / total_p * 100, 1)
        rejection_pct = round(d['rejected_period'] / total_p * 100, 1)
        kpi_body = (
            f"За отчётный период созданы и обработаны {total_p} документов. "
            f"Из них утверждено {d['approved_period']} ({approval_pct}%), "
            f"отклонено {d['rejected_period']} ({rejection_pct}%). "
        )
        if approval_pct >= 80:
            kpi_body += (
                "Высокий показатель утверждения свидетельствует о качественной подготовке документов "
                "и слаженной работе сотрудников на всех этапах согласования. "
            )
        elif approval_pct >= 50:
            kpi_body += (
                "Уровень утверждения находится в удовлетворительном диапазоне. "
                "Вместе с тем прослеживается потенциал для улучшения качества подготовки документов "
                "с целью сокращения количества отклонений и повторных итераций. "
            )
        else:
            kpi_body += (
                "Уровень утверждения ниже рекомендуемого порогового значения (50%), "
                "что может указывать на системные проблемы с качеством исходных документов, "
                "недостаточную компетентность исполнителей или нечёткость регламентов согласования. "
                "Рекомендуется провести детальный анализ причин отказов. "
            )
        if d['overdue_count'] > 0:
            overdue_pct = round(d['overdue_count'] / max(d['in_work_period'], 1) * 100, 1)
            kpi_body += (
                f"Зафиксировано {d['overdue_count']} активных просроченных документов "
                f"({overdue_pct}% от документов в работе). "
            )
            if d['overdue_count'] >= 10:
                kpi_body += (
                    "Значительное число просрочек требует немедленного управленческого внимания "
                    "и мер по нормализации исполнительской дисциплины. "
                )
        if d['avg_days'] is not None:
            kpi_body += (
                f"Среднее время обработки (от создания до утверждения) составило {d['avg_days']} "
                f"{'день' if d['avg_days'] == 1 else 'дня' if 2 <= d['avg_days'] <= 4 else 'дней'}. "
            )
            if d['avg_days'] <= 2:
                kpi_body += "Это свидетельствует об эффективной и быстрой обработке документов. "
            elif d['avg_days'] <= 7:
                kpi_body += "Скорость обработки находится в стандартном диапазоне для большинства организаций. "
            else:
                kpi_body += (
                    "Длительный цикл обработки может негативно сказываться на оперативности принятия решений. "
                    "Рекомендуется пересмотреть этапы маршрутизации документов. "
                )
        sections.append({'title': 'Анализ ключевых показателей эффективности (КПЭ)', 'body': kpi_body})

    # ── 3. Структурный анализ ─────────────────────────────────────────────────
    if d['by_status']:
        dominant = max(d['by_status'], key=lambda x: x['count'])
        dominant_lbl = d['STATUS_LABELS'].get(dominant['status'], dominant['status'])
        dominant_pct = round(dominant['count'] / (total_p or 1) * 100, 1)
        struct_body = (
            f"Анализ распределения документов по статусам показывает, что преобладающим является статус "
            f"«{dominant_lbl}» — {dominant['count']} документов ({dominant_pct}%). "
        )
        status_map = {i['status']: i['count'] for i in d['by_status']}
        draft_cnt = status_map.get('draft', 0)
        if draft_cnt > 0 and total_p > 0:
            draft_pct = round(draft_cnt / total_p * 100, 1)
            if draft_pct > 30:
                struct_body += (
                    f"Обращает на себя внимание высокая доля документов в статусе «Черновик» ({draft_pct}%), "
                    "что может указывать на незавершённые рабочие процессы или отсутствие стимулов "
                    "для своевременной подачи документов на согласование. "
                )
        if d['by_type']:
            top_type = d['by_type'][0]
            top_type_lbl = d['TYPE_LABELS'].get(top_type['template__type'], top_type['template__type'] or 'Прочие')
            struct_body += (
                f"По типовому составу лидирует категория «{top_type_lbl}» "
                f"({top_type['count']} ед., {round(top_type['count']/(total_p or 1)*100, 1)}%). "
            )
            if len(d['by_type']) > 1:
                struct_body += (
                    f"Всего в периоде зафиксировано {len(d['by_type'])} различных типов документов, "
                    "что свидетельствует о многопрофильном характере документооборота организации. "
                )
        sections.append({'title': 'Структурный анализ документооборота', 'body': struct_body})

    # ── 4. Процессы согласования ──────────────────────────────────────────────
    if d['wf_total'] > 0:
        wf_appr_pct = round(d['wf_approved'] / d['wf_total'] * 100, 1)
        wf_rej_pct  = round(d['wf_rejected'] / d['wf_total'] * 100, 1)
        wf_pend_pct = round(d['wf_pending']  / d['wf_total'] * 100, 1)
        wf_body = (
            f"В рамках маршрутов согласования за период принято {d['wf_total']} решений. "
            f"Из них одобрено {d['wf_approved']} ({wf_appr_pct}%), "
            f"отклонено {d['wf_rejected']} ({wf_rej_pct}%), "
            f"ожидают решения {d['wf_pending']} ({wf_pend_pct}%). "
        )
        if wf_appr_pct >= 75:
            wf_body += (
                "Высокий процент одобрения в маршрутах согласования свидетельствует о зрелости "
                "внутренних регламентов и высоком уровне взаимодействия между подразделениями. "
            )
        elif wf_appr_pct >= 50:
            wf_body += (
                "Процент одобрения в маршрутах согласования находится на среднем уровне. "
                "Анализ причин отклонений поможет выявить слабые звенья в регламентах. "
            )
        else:
            wf_body += (
                "Низкий процент одобрения в маршрутах согласования является тревожным сигналом. "
                "Возможные причины: несоответствие документов требованиям, нечёткие критерии одобрения, "
                "или неверно настроенные маршруты. Рекомендуется провести аудит шаблонов согласования. "
            )
        if d['wf_pending'] > 0:
            wf_body += (
                f"Наличие {d['wf_pending']} ожидающих решений требует внимания ответственных лиц "
                "во избежание накопления необработанных задач в очереди согласования. "
            )
        sections.append({'title': 'Анализ процессов согласования', 'body': wf_body})

    # ── 5. Анализ активности пользователей ───────────────────────────────────
    if d['by_creator'] or d['by_assignee']:
        user_body = ""
        if d['by_creator']:
            top_cr = d['by_creator'][0]
            fn = top_cr.get('created_by__first_name', '')
            ln = top_cr.get('created_by__last_name', '')
            un = top_cr.get('created_by__username', '')
            top_creator_name = f'{fn} {ln}'.strip() or un
            creator_pct = round(top_cr['count'] / (total_p or 1) * 100, 1)
            user_body += (
                f"В части активности по созданию документов лидирует сотрудник "
                f"{top_creator_name} — {top_cr['count']} документов ({creator_pct}% от общего объёма). "
            )
            if len(d['by_creator']) > 1:
                all_top_pct = round(sum(i['count'] for i in d['by_creator']) / (total_p or 1) * 100, 1)
                user_body += (
                    f"Топ-{len(d['by_creator'])} авторов в совокупности обеспечивают "
                    f"{all_top_pct}% всего документопотока периода. "
                )
                if creator_pct > 60:
                    user_body += (
                        "Высокая концентрация документопотока у единственного автора "
                        "может создавать риски «узкого места» в бизнес-процессах. "
                    )
        if d['by_assignee']:
            top_as = d['by_assignee'][0]
            fn = top_as.get('assigned_to__first_name', '')
            ln = top_as.get('assigned_to__last_name', '')
            un = top_as.get('assigned_to__username', '')
            top_assignee_name = f'{fn} {ln}'.strip() or un
            user_body += (
                f"Наиболее нагруженным исполнителем является {top_assignee_name} "
                f"({top_as['count']} документов назначено). "
            )
            if len(d['by_assignee']) >= 3:
                loads = [i['count'] for i in d['by_assignee']]
                if loads[0] > loads[-1] * 3:
                    user_body += (
                        "Распределение нагрузки среди исполнителей неравномерно: "
                        "отдельные сотрудники перегружены по сравнению с другими. "
                        "Рекомендуется балансировать назначения для более равномерного распределения задач. "
                    )
                else:
                    user_body += "Распределение нагрузки среди исполнителей достаточно равномерное. "
        sections.append({'title': 'Анализ активности участников процесса', 'body': user_body})

    # ── 6. Блок рисков / просроченные ────────────────────────────────────────
    overdue_cnt = d['overdue_count']
    if overdue_cnt > 0:
        risk_body = (
            f"По состоянию на дату формирования отчёта в системе зафиксировано "
            f"{overdue_cnt} {'просроченный документ' if overdue_cnt == 1 else 'просроченных документа' if 2 <= overdue_cnt <= 4 else 'просроченных документов'} "
            f"с нарушением установленных сроков исполнения. "
        )
        if overdue_cnt >= 15:
            risk_body += (
                "Высокий уровень просрочки несёт существенные операционные и репутационные риски. "
                "Рекомендованы: немедленное информирование ответственных исполнителей, "
                "назначение дополнительных ресурсов на приоритетные документы, "
                "а также системный пересмотр сроков и регламентов обработки. "
            )
        elif overdue_cnt >= 5:
            risk_body += (
                "Умеренный уровень просрочки требует оперативного вмешательства руководителей. "
                "Необходимо выявить причины задержек (загруженность исполнителей, "
                "нечёткость ТЗ, технические проблемы) и принять меры по устранению. "
            )
        else:
            risk_body += (
                "Незначительное число просрочек находится в пределах допустимого уровня. "
                "Тем не менее рекомендуется лично уведомить ответственных исполнителей "
                "и установить чёткие сроки завершения обработки. "
            )
        if d['overdue_list']:
            oldest = None
            for odoc in d['overdue_list']:
                if odoc.deadline:
                    if oldest is None or odoc.deadline < oldest.deadline:
                        oldest = odoc
            if oldest:
                risk_body += (
                    f"Наиболее длительная просрочка зафиксирована по документу "
                    f"«{oldest.title[:50]}» (рег. № {oldest.registry_number or 'б/н'}, "
                    f"срок — {oldest.deadline.strftime('%d.%m.%Y')}). "
                )
        sections.append({'title': 'Оценка рисков и просроченные документы', 'body': risk_body})
    else:
        sections.append({
            'title': 'Оценка рисков и просроченные документы',
            'body': (
                "На дату формирования отчёта просроченных документов не выявлено. "
                "Исполнительская дисциплина соответствует установленным стандартам, "
                "все активные документы обрабатываются в установленные сроки. "
                "Рекомендуется поддерживать текущий уровень дисциплины и проводить "
                "регулярный мониторинг сроков исполнения для предотвращения просрочек."
            )
        })

    # ── 7. Выводы и рекомендации ──────────────────────────────────────────────
    concl_body = f"По итогам анализа данных за период «{d['period_label']}» можно сделать следующие выводы:\n"
    points = []
    if total_p > 0:
        approval_pct = round(d['approved_period'] / total_p * 100, 1)
        points.append(
            f"уровень утверждения документов составил {approval_pct}% "
            f"({'выше нормы' if approval_pct >= 70 else 'требует улучшения'})"
        )
    if d['overdue_count'] == 0:
        points.append("просроченных документов не зафиксировано — исполнительская дисциплина в норме")
    else:
        points.append(f"выявлено {d['overdue_count']} просроченных документов, требующих немедленного внимания")
    if d['avg_days'] is not None:
        points.append(f"среднее время обработки документа — {d['avg_days']} дн.")
    if d['wf_total'] > 0:
        wf_appr_pct = round(d['wf_approved'] / d['wf_total'] * 100, 1)
        points.append(f"в маршрутах согласования одобрено {wf_appr_pct}% решений")

    concl_body = (
        f"По итогам анализа данных за период «{d['period_label']}» сформированы следующие выводы: "
        + "; ".join(points) + ". "
    )
    concl_recs = []
    if d['overdue_count'] > 0:
        concl_recs.append("усилить контроль исполнительской дисциплины по просроченным документам")
    if total_p > 0 and d['approved_period'] / total_p < 0.5:
        concl_recs.append("провести работу по повышению качества подготовки документов")
    if d['avg_days'] is not None and d['avg_days'] > 7:
        concl_recs.append("оптимизировать маршруты согласования для сокращения цикла обработки")
    if d['wf_pending'] > 0:
        concl_recs.append("обеспечить закрытие ожидающих согласований в кратчайшие сроки")
    if concl_recs:
        concl_body += "Рекомендации: " + "; ".join(concl_recs) + ". "
    concl_body += (
        "Представленные показатели и выводы рекомендуется использовать при принятии управленческих решений, "
        "планировании ресурсов и аудите внутренних регламентов документооборота."
    )
    sections.append({'title': 'Заключение и рекомендации', 'body': concl_body})

    return sections


# ─────────────────────────────────────────────────────────────────────────────
# Report export: PDF
# ─────────────────────────────────────────────────────────────────────────────
@login_required
@manager_required
def report_export_pdf(request):
    """Экспорт аналитического отчёта в PDF"""
    from io import BytesIO
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    d = _report_get_data(request)

    # Try to register a Cyrillic-capable font
    font_name = 'Helvetica'
    _fonts_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), '..', 'static', 'fonts'))
    for _reg, _bold in [
        (os.path.join(_fonts_dir, 'DejaVuSans.ttf'), os.path.join(_fonts_dir, 'DejaVuSans-Bold.ttf')),
        ('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'),
        ('C:/Windows/Fonts/arial.ttf', 'C:/Windows/Fonts/arial.ttf'),
        ('C:/Windows/Fonts/Arial.ttf', 'C:/Windows/Fonts/Arial.ttf'),
    ]:
        if os.path.exists(_reg):
            try:
                pdfmetrics.registerFont(TTFont('CyrFont', _reg))
                _b = _bold if os.path.exists(_bold) else _reg
                pdfmetrics.registerFont(TTFont('CyrFont-Bold', _b))
                font_name = 'CyrFont'
            except Exception:
                pass
            break

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('RpTitle', fontName=font_name, fontSize=18, textColor=colors.HexColor('#312E81'),
                                  spaceAfter=6, spaceBefore=0, leading=22)
    sub_style   = ParagraphStyle('RpSub',   fontName=font_name, fontSize=10, textColor=colors.HexColor('#6B7280'),
                                  spaceAfter=16)
    h2_style    = ParagraphStyle('RpH2',    fontName=font_name, fontSize=12, textColor=colors.HexColor('#1E1B4B'),
                                  spaceBefore=14, spaceAfter=4, leading=16)
    body_style  = ParagraphStyle('RpBody',  fontName=font_name, fontSize=10, textColor=colors.HexColor('#374151'),
                                  spaceBefore=2, spaceAfter=10, leading=15)
    label_style = ParagraphStyle('RpLabel', fontName=font_name, fontSize=9, textColor=colors.HexColor('#6B7280'),
                                  spaceBefore=0, spaceAfter=4, leading=13)

    INDIGO  = colors.HexColor('#4338CA')
    INDIGO_L= colors.HexColor('#E0E7FF')
    GRAY    = colors.HexColor('#F9FAFB')
    BORDER  = colors.HexColor('#E5E7EB')

    def make_table(headers, rows, col_widths=None):
        data = [headers] + rows
        t = Table(data, colWidths=col_widths, repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND',  (0, 0), (-1, 0),  INDIGO),
            ('TEXTCOLOR',   (0, 0), (-1, 0),  colors.white),
            ('FONTNAME',    (0, 0), (-1, 0),  font_name),
            ('FONTSIZE',    (0, 0), (-1, 0),  9),
            ('FONTNAME',    (0, 1), (-1, -1), font_name),
            ('FONTSIZE',    (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, GRAY]),
            ('GRID',        (0, 0), (-1, -1), 0.4, BORDER),
            ('ALIGN',       (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN',      (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING',(0, 0), (-1, -1), 6),
            ('TOPPADDING',  (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
        ]))
        return t

    def add_narrative(story, sections_list, section_title):
        """Найти секцию по заголовку и добавить её текст в story."""
        for sec in sections_list:
            if sec['title'] == section_title:
                story.append(Paragraph(sec['body'], body_style))
                break

    narrative = _report_build_narrative(d)

    story = []
    story.append(Paragraph('Аналитический отчёт', title_style))
    story.append(Paragraph(
        f'Период: {d["period_label"]}   |   Сформирован: {d["today"].strftime("%d.%m.%Y %H:%M")}',
        sub_style))
    story.append(HRFlowable(width='100%', thickness=1, color=INDIGO_L))
    story.append(Spacer(1, 8))

    # ── Введение ──────────────────────────────────────────────────────────────
    story.append(Paragraph('1. Введение', h2_style))
    add_narrative(story, narrative, 'Введение')

    # ── КПЭ ──────────────────────────────────────────────────────────────────
    story.append(Paragraph('2. Ключевые показатели эффективности', h2_style))
    story.append(Paragraph('Таблица 1. Сводные показатели за период', label_style))
    kpi_rows = [
        ['Всего документов в системе',     str(d['total_all'])],
        ['Документов за период',           str(d['total_period'])],
        ['Утверждено за период',           str(d['approved_period'])],
        ['Отклонено за период',            str(d['rejected_period'])],
        ['В работе за период',             str(d['in_work_period'])],
        ['На рассмотрении за период',      str(d['in_review_period'])],
        ['Просрочено (активных)',          str(d['overdue_count'])],
        ['Ср. время обработки (дней)',     str(d['avg_days']) if d['avg_days'] is not None else '—'],
    ]
    story.append(make_table(['Показатель', 'Значение'], kpi_rows, [11*cm, 4*cm]))
    story.append(Spacer(1, 6))
    add_narrative(story, narrative, 'Анализ ключевых показателей эффективности (КПЭ)')

    # ── Согласование ─────────────────────────────────────────────────────────
    if d['wf_total']:
        story.append(Paragraph('3. Процессы согласования', h2_style))
        story.append(Paragraph('Таблица 2. Результаты согласования (WorkflowApproval)', label_style))
        wf_rows = [
            ['Утверждено', str(d['wf_approved'])],
            ['Отклонено',  str(d['wf_rejected'])],
            ['Ожидает',    str(d['wf_pending'])],
            ['Всего',      str(d['wf_total'])],
        ]
        story.append(make_table(['Результат', 'Количество'], wf_rows, [9*cm, 4*cm]))
        story.append(Spacer(1, 6))
        add_narrative(story, narrative, 'Анализ процессов согласования')

    # ── Структура ─────────────────────────────────────────────────────────────
    total_p = d['total_period'] or 1
    sec_num = 4 if d['wf_total'] else 3
    story.append(Paragraph(f'{sec_num}. Структура документооборота', h2_style))
    story.append(Paragraph(f'Таблица {sec_num}а. Разбивка по статусам', label_style))
    st_rows = [[d['STATUS_LABELS'].get(i['status'], i['status']), str(i['count']),
                f"{round(i['count']/total_p*100,1)}%"] for i in d['by_status']]
    story.append(make_table(['Статус', 'Кол-во', 'Доля'], st_rows, [9*cm, 3*cm, 3*cm]))
    story.append(Spacer(1, 6))

    if d['by_type']:
        story.append(Paragraph(f'Таблица {sec_num}б. Разбивка по типам документов', label_style))
        tp_rows = [[d['TYPE_LABELS'].get(i['template__type'], i['template__type'] or '—'),
                    str(i['count']), f"{round(i['count']/total_p*100,1)}%"] for i in d['by_type']]
        story.append(make_table(['Тип', 'Кол-во', 'Доля'], tp_rows, [9*cm, 3*cm, 3*cm]))
        story.append(Spacer(1, 6))

    add_narrative(story, narrative, 'Структурный анализ документооборота')

    # ── Активность пользователей ──────────────────────────────────────────────
    sec_num2 = sec_num + 1
    if d['by_creator'] or d['by_assignee']:
        story.append(Paragraph(f'{sec_num2}. Активность участников процесса', h2_style))
    if d['by_creator']:
        story.append(Paragraph(f'Таблица {sec_num2}а. Топ авторов документов', label_style))
        cr_rows = []
        for idx, item in enumerate(d['by_creator'], 1):
            fn = item.get('created_by__first_name', '')
            ln = item.get('created_by__last_name', '')
            un = item.get('created_by__username', '')
            name = f'{fn} {ln}'.strip() or un
            cr_rows.append([str(idx), name, str(item['count'])])
        story.append(make_table(['#', 'Сотрудник', 'Создано'], cr_rows, [1.5*cm, 9.5*cm, 4*cm]))
        story.append(Spacer(1, 6))

    if d['by_assignee']:
        story.append(Paragraph(f'Таблица {sec_num2}б. Топ исполнителей', label_style))
        as_rows = []
        for idx, item in enumerate(d['by_assignee'], 1):
            fn = item.get('assigned_to__first_name', '')
            ln = item.get('assigned_to__last_name', '')
            un = item.get('assigned_to__username', '')
            name = f'{fn} {ln}'.strip() or un
            as_rows.append([str(idx), name, str(item['count'])])
        story.append(make_table(['#', 'Сотрудник', 'Назначено'], as_rows, [1.5*cm, 9.5*cm, 4*cm]))
        story.append(Spacer(1, 6))

    if d['by_creator'] or d['by_assignee']:
        add_narrative(story, narrative, 'Анализ активности участников процесса')

    # ── Просроченные / риски ──────────────────────────────────────────────────
    sec_num3 = sec_num2 + 1
    story.append(Paragraph(f'{sec_num3}. Оценка рисков и просроченные документы', h2_style))
    add_narrative(story, narrative, 'Оценка рисков и просроченные документы')
    if d['overdue_list']:
        story.append(Paragraph(f'Таблица {sec_num3}. Список просроченных документов', label_style))
        ov_rows = []
        for odoc in d['overdue_list']:
            assignee = odoc.assigned_to.get_full_name() if odoc.assigned_to else '—'
            deadline = odoc.deadline.strftime('%d.%m.%Y') if odoc.deadline else '—'
            ov_rows.append([odoc.registry_number or '—', odoc.title[:50], deadline, assignee])
        story.append(make_table(['Рег. №', 'Название', 'Срок', 'Ответственный'], ov_rows,
                                [3*cm, 7*cm, 3*cm, 4*cm]))
        story.append(Spacer(1, 6))

    # ── Заключение ────────────────────────────────────────────────────────────
    sec_num4 = sec_num3 + 1
    story.append(HRFlowable(width='100%', thickness=0.5, color=INDIGO_L))
    story.append(Spacer(1, 4))
    story.append(Paragraph(f'{sec_num4}. Заключение и рекомендации', h2_style))
    add_narrative(story, narrative, 'Заключение и рекомендации')

    doc.build(story)
    buf.seek(0)
    fname = f'report_{d["today"].strftime("%Y%m%d")}.pdf'
    resp = FileResponse(buf, content_type='application/pdf')
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    return resp


# ─────────────────────────────────────────────────────────────────────────────
# Report export: DOCX
# ─────────────────────────────────────────────────────────────────────────────
@login_required
@manager_required
def report_export_docx(request):
    """Экспорт аналитического отчёта в Word (.docx)"""
    from io import BytesIO
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    d = _report_get_data(request)
    doc = DocxDoc()

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    INDIGO = RGBColor(0x43, 0x38, 0xCA)
    DARK   = RGBColor(0x11, 0x18, 0x27)
    GRAY   = RGBColor(0x6B, 0x72, 0x80)

    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def add_section_heading(doc, text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = INDIGO

    def add_table_label(doc, text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        for run in p.runs:
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = GRAY

    def add_body_text(doc, text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(8)
        for run in p.runs:
            run.font.size = Pt(10)

    def add_kv_table(doc, headers, rows):
        tbl = doc.add_table(rows=1, cols=len(headers))
        tbl.style = 'Table Grid'
        hdr_row = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            shade_cell(cell, '4338CA')
            run = cell.paragraphs[0].add_run(h)
            run.bold = True; run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for r_vals in rows:
            row = tbl.add_row()
            for i, v in enumerate(r_vals):
                row.cells[i].text = str(v)
                row.cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        return tbl

    def add_narrative_section(doc, sections_list, section_title):
        for sec in sections_list:
            if sec['title'] == section_title:
                add_body_text(doc, sec['body'])
                break

    narrative = _report_build_narrative(d)

    # ── Title block ──────────────────────────────────────────
    title_p = doc.add_heading('', level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run('Аналитический отчёт')
    run.font.size = Pt(18); run.font.color.rgb = INDIGO

    sub = doc.add_paragraph(f'Период: {d["period_label"]}')
    sub.runs[0].font.color.rgb = GRAY; sub.runs[0].font.size = Pt(10)

    sub2 = doc.add_paragraph(f'Сформирован: {d["today"].strftime("%d.%m.%Y %H:%M")}')
    sub2.runs[0].font.color.rgb = GRAY; sub2.runs[0].font.size = Pt(10)
    sub2.paragraph_format.space_after = Pt(10)

    # ── 1. Введение ─────────────────────────────────────────
    add_section_heading(doc, '1. Введение')
    add_narrative_section(doc, narrative, 'Введение')

    # ── 2. Ключевые показатели ───────────────────────────────
    add_section_heading(doc, '2. Ключевые показатели эффективности')
    add_table_label(doc, 'Таблица 1. Сводные показатели за период')
    kpi_rows = [
        ('Всего документов в системе',    d['total_all']),
        ('Документов за период',          d['total_period']),
        ('Утверждено за период',          d['approved_period']),
        ('Отклонено за период',           d['rejected_period']),
        ('В работе за период',             d['in_work_period']),
        ('На рассмотрении за период',      d['in_review_period']),
        ('Просрочено (активных)',         d['overdue_count']),
        ('Ср. время обработки (дней)',    d['avg_days'] if d['avg_days'] is not None else '—'),
    ]
    add_kv_table(doc, ['Показатель', 'Значение'], kpi_rows)
    doc.add_paragraph()
    add_narrative_section(doc, narrative, 'Анализ ключевых показателей эффективности (КПЭ)')

    # ── 3. Согласование ─────────────────────────────────────
    if d['wf_total']:
        add_section_heading(doc, '3. Процессы согласования')
        add_table_label(doc, 'Таблица 2. Результаты согласования (WorkflowApproval)')
        add_kv_table(doc, ['Результат', 'Количество'], [
            ('Утверждено', d['wf_approved']),
            ('Отклонено',  d['wf_rejected']),
            ('Ожидает',    d['wf_pending']),
            ('Всего',      d['wf_total']),
        ])
        doc.add_paragraph()
        add_narrative_section(doc, narrative, 'Анализ процессов согласования')

    # ── 4. Структура ──────────────────────────────────────────
    total_p = d['total_period'] or 1
    sec_num = 4 if d['wf_total'] else 3
    add_section_heading(doc, f'{sec_num}. Структура документооборота')
    add_table_label(doc, f'Таблица {sec_num}а. Разбивка по статусам')
    st_rows = [(d['STATUS_LABELS'].get(i['status'], i['status']), i['count'],
                f"{round(i['count']/total_p*100,1)}%") for i in d['by_status']]
    add_kv_table(doc, ['Статус', 'Количество', 'Доля'], st_rows)
    doc.add_paragraph()

    if d['by_type']:
        add_table_label(doc, f'Таблица {sec_num}б. Разбивка по типам документов')
        tp_rows = [(d['TYPE_LABELS'].get(i['template__type'], i['template__type'] or '—'),
                    i['count'], f"{round(i['count']/total_p*100,1)}%") for i in d['by_type']]
        add_kv_table(doc, ['Тип документа', 'Количество', 'Доля'], tp_rows)
        doc.add_paragraph()

    add_narrative_section(doc, narrative, 'Структурный анализ документооборота')

    # ── 5. Активность пользователей ──────────────────────────
    sec_num2 = sec_num + 1
    if d['by_creator'] or d['by_assignee']:
        add_section_heading(doc, f'{sec_num2}. Активность участников процесса')
    if d['by_creator']:
        add_table_label(doc, f'Таблица {sec_num2}а. Топ авторов документов')
        cr_rows = []
        for idx, item in enumerate(d['by_creator'], 1):
            fn = item.get('created_by__first_name', '')
            ln = item.get('created_by__last_name', '')
            un = item.get('created_by__username', '')
            cr_rows.append((idx, f'{fn} {ln}'.strip() or un, item['count']))
        add_kv_table(doc, ['#', 'Сотрудник', 'Создано'], cr_rows)
        doc.add_paragraph()

    if d['by_assignee']:
        add_table_label(doc, f'Таблица {sec_num2}б. Топ исполнителей')
        as_rows = []
        for idx, item in enumerate(d['by_assignee'], 1):
            fn = item.get('assigned_to__first_name', '')
            ln = item.get('assigned_to__last_name', '')
            un = item.get('assigned_to__username', '')
            as_rows.append((idx, f'{fn} {ln}'.strip() or un, item['count']))
        add_kv_table(doc, ['#', 'Сотрудник', 'Назначено'], as_rows)
        doc.add_paragraph()

    if d['by_creator'] or d['by_assignee']:
        add_narrative_section(doc, narrative, 'Анализ активности участников процесса')

    # ── 6. Риски / просроченные ──────────────────────────────
    sec_num3 = sec_num2 + 1
    add_section_heading(doc, f'{sec_num3}. Оценка рисков и просроченные документы')
    add_narrative_section(doc, narrative, 'Оценка рисков и просроченные документы')
    if d['overdue_list']:
        add_table_label(doc, f'Таблица {sec_num3}. Список просроченных документов')
        ov_rows = []
        for doc_obj in d['overdue_list']:
            assignee = doc_obj.assigned_to.get_full_name() if doc_obj.assigned_to else '—'
            deadline = doc_obj.deadline.strftime('%d.%m.%Y') if doc_obj.deadline else '—'
            status   = d['STATUS_LABELS'].get(doc_obj.status, doc_obj.status)
            ov_rows.append((doc_obj.registry_number or '—', doc_obj.title[:60], deadline, assignee, status))
        add_kv_table(doc, ['Рег. №', 'Название', 'Срок', 'Ответственный', 'Статус'], ov_rows)
        doc.add_paragraph()

    # ── 7. Заключение ────────────────────────────────────────
    sec_num4 = sec_num3 + 1
    add_section_heading(doc, f'{sec_num4}. Заключение и рекомендации')
    add_narrative_section(doc, narrative, 'Заключение и рекомендации')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    fname = f'report_{d["today"].strftime("%Y%m%d")}.docx'
    resp = HttpResponse(buf, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    resp['Content-Disposition'] = f'attachment; filename="{fname}"'
    return resp


@login_required
def notifications_list(request):
    """Список уведомлений пользователя"""
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
    return render(request, 'documents/notifications_modern.html', {'notifications': notifications})


@login_required
def mark_notification_read(request, pk):
    """Отметить уведомление как прочитанное"""
    notification = get_object_or_404(Notification, pk=pk, user=request.user)
    notification.is_read = True
    notification.save()
    return redirect('documents:notifications')


@login_required
def mark_all_notifications_read(request):
    """Отметить все уведомления как прочитанные"""
    Notification.objects.filter(user=request.user, is_read=False).update(is_read=True)
    messages.success(request, 'Все уведомления отмечены как прочитанные!')
    return redirect('documents:notifications')


# ============================================
# Chat Views
# ============================================

@login_required
def chat_users_list(request):
    """Список пользователей для чата (API)"""
    from django.contrib.auth.models import User
    
    users = User.objects.exclude(id=request.user.id).select_related('profile')
    users_data = []
    
    for user in users:
        # Количество непрочитанных сообщений от этого пользователя
        unread_count = ChatMessage.objects.filter(
            sender=user,
            recipient=request.user,
            is_read=False
        ).count()
        
        users_data.append({
            'id': user.id,
            'username': user.username,
            'full_name': user.get_full_name() or user.username,
            'role': user.profile.get_role_display() if hasattr(user, 'profile') else 'Пользователь',
            'avatar_url': user.profile.avatar.url if hasattr(user, 'profile') and user.profile.avatar else None,
            'unread_count': unread_count
        })
    
    return JsonResponse({'users': users_data})


@login_required
def chat_messages(request, user_id):
    """Получить историю сообщений с пользователем (API)"""
    from django.contrib.auth.models import User
    
    other_user = get_object_or_404(User, id=user_id)
    
    # Получаем все сообщения между текущим пользователем и выбранным
    messages_qs = ChatMessage.objects.filter(
        Q(sender=request.user, recipient=other_user) |
        Q(sender=other_user, recipient=request.user)
    ).select_related('sender', 'recipient').order_by('created_at')
    
    # Отмечаем полученные сообщения как прочитанные
    ChatMessage.objects.filter(
        sender=other_user,
        recipient=request.user,
        is_read=False
    ).update(is_read=True)
    
    messages_data = []
    for msg in messages_qs:
        messages_data.append({
            'id': msg.id,
            'sender_id': msg.sender.id,
            'sender_name': msg.sender.get_full_name() or msg.sender.username,
            'message': msg.message,
            'created_at': msg.created_at.strftime('%d.%m.%Y %H:%M'),
            'is_own': msg.sender == request.user
        })
    
    return JsonResponse({
        'messages': messages_data,
        'other_user': {
            'id': other_user.id,
            'full_name': other_user.get_full_name() or other_user.username
        }
    })


@login_required
def chat_send_message(request):
    """Отправить сообщение (API)"""
    if request.method == 'POST':
        from django.contrib.auth.models import User
        
        try:
            data = json.loads(request.body)
            recipient_id = data.get('recipient_id')
            message_text = data.get('message', '').strip()
            
            if not message_text:
                return JsonResponse({'error': 'Сообщение не может быть пустым'}, status=400)
            
            recipient = get_object_or_404(User, id=recipient_id)
            
            # Создаем сообщение
            message = ChatMessage.objects.create(
                sender=request.user,
                recipient=recipient,
                message=message_text
            )
            
            # Создаем уведомление для получателя
            Notification.objects.create(
                user=recipient,
                message=f"Новое сообщение от {request.user.get_full_name() or request.user.username}: {message_text[:50]}{'...' if len(message_text) > 50 else ''}"
            )
            # Отправка email
            send_chat_message_email(request.user, recipient, message_text)
            
            return JsonResponse({
                'success': True,
                'message': {
                    'id': message.id,
                    'sender_id': message.sender.id,
                    'sender_name': message.sender.get_full_name() or message.sender.username,
                    'message': message.message,
                    'created_at': message.created_at.strftime('%d.%m.%Y %H:%M'),
                    'is_own': True
                }
            })
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Метод не разрешен'}, status=405)


@login_required
def chat_unread_count(request):
    """Получить количество непрочитанных сообщений (API)"""
    unread_count = ChatMessage.objects.filter(
        recipient=request.user,
        is_read=False
    ).count()
    
    return JsonResponse({'unread_count': unread_count})


# ============================================
# CLERK FUNCTIONALITY - QUICK TEMPLATES
# ============================================

@login_required
def quick_templates(request):
    """Быстрые шаблоны — доступны всем авторизованным"""
    from .models import QuickTemplate

    templates = QuickTemplate.objects.filter(is_active=True)
    
    return render(request, 'documents/quick_templates.html', {
        'templates': templates
    })


@login_required
def create_from_quick_template(request, template_id):
    """Создать документ из быстрого шаблона — доступно всем авторизованным"""
    from .models import QuickTemplate

    template = get_object_or_404(QuickTemplate, id=template_id)
    template.usage_count += 1
    template.save()
    
    if request.method == 'POST':
        title = request.POST.get('title')
        description = request.POST.get('description')
        recipients_ids = request.POST.getlist('recipients')
        
        # Сохраняем тип документа в метаданных
        document = Document.objects.create(
            title=title,
            content=template.content_template,
            created_by=request.user,
            status='draft',
            metadata={'document_type': template.document_type, 'description': description}
        )
        
        from django.contrib.auth.models import User
        recipients = User.objects.filter(id__in=recipients_ids)
        if recipients.exists():
            document.assigned_to = recipients.first()
            document.save()
        
        messages.success(request, 'Документ создан из шаблона!')
        return redirect('documents:document_detail', pk=document.id)
    
    from django.contrib.auth.models import User
    users = User.objects.exclude(id=request.user.id)
    
    return render(request, 'documents/create_from_quick_template.html', {
        'template': template,
        'users': users
    })


# ============================================
# MANAGER FUNCTIONALITY - APPROVALS & STATS
# ============================================

@login_required
def manager_approvals(request):
    """Одобрения для руководителей"""
    from .models import WorkflowApproval

    if request.user.profile.role not in ['manager', 'admin']:
        return redirect('documents:dashboard')

    pending_approvals = WorkflowApproval.objects.filter(
        approver=request.user,
        decision='pending'
    ).select_related('workflow_step__document')

    # Documents in in_review status that have no pending workflow approval
    # (i.e. they need direct manager action)
    in_review_docs = Document.objects.filter(
        status='in_review'
    ).exclude(
        workflow_steps__approvals__decision='pending'
    ).select_related('created_by', 'assigned_to', 'template').order_by('-created_at')

    completed_approvals = WorkflowApproval.objects.filter(
        approver=request.user
    ).exclude(decision='pending').select_related('workflow_step__document').order_by('-decision_date')[:20]

    return render(request, 'documents/manager_approvals_modern.html', {
        'pending_approvals': pending_approvals,
        'in_review_docs': in_review_docs,
        'completed_approvals': completed_approvals,
    })


@login_required
def process_approval(request, approval_id):
    """Обработка одобрения"""
    from .models import WorkflowApproval
    from django.utils import timezone
    
    if request.user.profile.role not in ['manager', 'admin']:
        return redirect('documents:dashboard')
    
    approval = get_object_or_404(WorkflowApproval, id=approval_id, approver=request.user)
    
    if request.method == 'POST':
        decision = request.POST.get('decision')
        comments = request.POST.get('comments', '')
        
        approval.decision = decision
        approval.comments = comments
        approval.decision_date = timezone.now()
        approval.save()
        
        # Обновляем статус шага workflow
        if decision == 'approved':
            approval.workflow_step.status = 'approved'
            approval.workflow_step.save()
            messages.success(request, 'Документ одобрен!')
        elif decision == 'rejected':
            approval.workflow_step.status = 'rejected'
            approval.workflow_step.save()
            messages.warning(request, 'Документ отклонен!')
        
        return redirect('documents:manager_approvals')
    
    return render(request, 'documents/process_approval.html', {
        'approval': approval
    })


@login_required
def chat_page(request, active_user_id=None):
    """Страница чата — полноэкранный чат с выбором собеседника"""
    from django.contrib.auth.models import User as AuthUser

    # Все пользователи, кроме текущего
    users = AuthUser.objects.exclude(id=request.user.id).select_related('profile').order_by(
        'last_name', 'first_name'
    )

    # Обогащаем список: считаем непрочитанные для каждого
    users_data = []
    for u in users:
        unread = ChatMessage.objects.filter(
            sender=u, recipient=request.user, is_read=False
        ).count()
        users_data.append({
            'user': u,
            'unread': unread,
        })

    # Активный собеседник
    active_user = None
    active_messages = []
    if active_user_id:
        active_user = get_object_or_404(AuthUser, pk=active_user_id)
        # Пометить входящие как прочитанные
        ChatMessage.objects.filter(
            sender=active_user, recipient=request.user, is_read=False
        ).update(is_read=True)
        # Загрузить историю
        active_messages = ChatMessage.objects.filter(
            Q(sender=request.user, recipient=active_user) |
            Q(sender=active_user, recipient=request.user)
        ).select_related('sender').order_by('created_at')

    # Обработка отправки сообщения
    if request.method == 'POST' and active_user:
        text = request.POST.get('message', '').strip()
        if text:
            ChatMessage.objects.create(
                sender=request.user,
                recipient=active_user,
                message=text
            )
            Notification.objects.create(
                user=active_user,
                message=f"Новое сообщение от {request.user.get_full_name() or request.user.username}: {text[:50]}{'...' if len(text) > 50 else ''}"
            )
        return redirect('documents:chat_with_user', active_user_id=active_user_id)

    return render(request, 'documents/chat_modern.html', {
        'users_data': users_data,
        'active_user': active_user,
        'active_messages': active_messages,
    })
