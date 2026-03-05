"""
Инициализация production БД: создаёт пользователей и тестовые данные при старте.
Идемпотентна — не создаёт дубли при повторных запусках.
"""
import random
from io import BytesIO
from datetime import timedelta

from django.core.management.base import BaseCommand
from django.contrib.auth.models import User
from django.core.files.base import ContentFile
from django.utils import timezone

from accounts.models import UserProfile
from documents.models import Document, DocumentTemplate


# ─── Обязательные пользователи ────────────────────────────────────────────────
REQUIRED_USERS = [
    {
        'username': 'admin',
        'password': 'admin',
        'first_name': 'Администратор',
        'last_name': 'Системы',
        'email': 'admin@example.com',
        'is_superuser': True,
        'is_staff': True,
        'role': 'admin',
        'position': 'Системный администратор',
        'department': 'IT отдел',
    },
    {
        'username': 'александров.алексей45',
        'password': 'Hfqcel779',
        'first_name': 'Алексей',
        'last_name': 'Александров',
        'email': 'alexey45@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'clerk',
        'position': 'Делопроизводитель',
        'department': 'Канцелярия',
    },
    {
        'username': 'александров.виктор47',
        'password': 'Hfqcel779',
        'first_name': 'Виктор',
        'last_name': 'Александров',
        'email': 'viktor47@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'manager',
        'position': 'Руководитель',
        'department': 'Управление',
    },
    {
        'username': 'александров.игорь10',
        'password': 'Hfqcel779',
        'first_name': 'Игорь',
        'last_name': 'Александров',
        'email': 'igor10@example.com',
        'is_superuser': False, 'is_staff': False,
        'role': 'employee',
        'position': 'Сотрудник',
        'department': 'Отдел исполнения',
    },
]

# ─── Дополнительные тестовые пользователи ─────────────────────────────────────
EXTRA_USERS = [
    ('Иванов', 'Дмитрий', 'manager', 'Руководитель отдела', 'Финансовый отдел'),
    ('Петров', 'Сергей', 'clerk', 'Старший делопроизводитель', 'Канцелярия'),
    ('Смирнова', 'Ольга', 'employee', 'Специалист', 'Бухгалтерия'),
    ('Козлов', 'Андрей', 'employee', 'Инженер', 'IT отдел'),
    ('Новикова', 'Елена', 'employee', 'Аналитик', 'Отдел продаж'),
    ('Федоров', 'Максим', 'manager', 'Начальник отдела', 'Отдел кадров'),
    ('Волкова', 'Наталья', 'clerk', 'Делопроизводитель', 'Юридический отдел'),
    ('Морозов', 'Павел', 'employee', 'Менеджер', 'Маркетинг'),
    ('Соколова', 'Мария', 'employee', 'Консультант', 'Отдел продаж'),
    ('Лебедев', 'Роман', 'employee', 'Специалист по закупкам', 'Закупки'),
    ('Захарова', 'Татьяна', 'employee', 'Бухгалтер', 'Бухгалтерия'),
    ('Егоров', 'Николай', 'manager', 'Директор департамента', 'Производство'),
    ('Кузнецова', 'Ирина', 'employee', 'HR-специалист', 'Отдел кадров'),
    ('Васильев', 'Константин', 'employee', 'Юрист', 'Юридический отдел'),
    ('Михайлова', 'Анна', 'clerk', 'Офис-менеджер', 'Административный отдел'),
]

# ─── Шаблоны ──────────────────────────────────────────────────────────────────
TEMPLATES_DATA = [
    ('Приказ о приёме на работу', 'order', 'docx',
     'Стандартный приказ о приёме сотрудника на работу'),
    ('Приказ об отпуске', 'order', 'docx',
     'Приказ о предоставлении ежегодного оплачиваемого отпуска'),
    ('Приказ о командировке', 'order', 'docx',
     'Приказ о направлении сотрудника в командировку'),
    ('Приказ о премировании', 'order', 'xlsx',
     'Приказ о выплате премии сотрудникам'),
    ('Трудовой договор', 'contract', 'docx',
     'Стандартный трудовой договор с сотрудником'),
    ('Договор поставки', 'contract', 'docx',
     'Договор на поставку товаров и услуг'),
    ('Договор подряда', 'contract', 'docx',
     'Договор подряда с подрядчиком'),
    ('Акт приёма-передачи', 'act', 'docx',
     'Акт приёма-передачи документов или имущества'),
    ('Акт выполненных работ', 'act', 'xlsx',
     'Акт о выполнении работ по договору'),
    ('Служебная записка', 'memo', 'docx',
     'Служебная записка для внутренней переписки'),
    ('Финансовый отчёт', 'report', 'xlsx',
     'Финансовый отчёт о доходах и расходах'),
    ('Квартальный отчёт', 'report', 'docx',
     'Отчёт о деятельности подразделения за квартал'),
    ('Письмо контрагенту', 'letter', 'docx',
     'Официальное письмо внешнему контрагенту'),
    ('Заявление на отпуск', 'application', 'docx',
     'Заявление сотрудника на предоставление отпуска'),
    ('Смета расходов', 'other', 'xlsx',
     'Смета расходов по проекту или мероприятию'),
]


# ─── Генераторы файлов ────────────────────────────────────────────────────────

def _make_docx(title, body_text):
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph()
    doc.add_heading('Плейсхолдеры', level=2)
    for ph in ['{{ФИО}}', '{{должность}}', '{{отдел}}', '{{дата}}', '{{номер}}']:
        doc.add_paragraph(ph, style='List Bullet')

    doc.add_paragraph()
    doc.add_heading('Раздел 1. Основные сведения', level=2)
    doc.add_paragraph('ФИО исполнителя: {{ФИО}}')
    doc.add_paragraph('Должность: {{должность}}')
    doc.add_paragraph('Отдел: {{отдел}}')

    doc.add_heading('Раздел 2. Реквизиты', level=2)
    doc.add_paragraph('Дата составления: {{дата}}')
    doc.add_paragraph('Номер документа: {{номер}}')

    doc.add_heading('Раздел 3. Содержание', level=2)
    doc.add_paragraph(body_text)
    doc.add_paragraph('Подпись: _______________')

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _make_xlsx(title):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = title[:31]

    headers = ['Параметр', 'Плейсхолдер', 'Значение', 'Примечание']
    rows = [
        ('ФИО', '{{ФИО}}', '', 'Заполнить'),
        ('Должность', '{{должность}}', '', 'Заполнить'),
        ('Отдел', '{{отдел}}', '', 'Заполнить'),
        ('Дата', '{{дата}}', '', 'Заполнить'),
        ('Номер', '{{номер}}', '', 'Заполнить'),
        ('Сумма', '{{сумма}}', '', 'Если применимо'),
        ('Период', '{{период}}', '', 'Если применимо'),
    ]

    ws.merge_cells('A1:D1')
    tc = ws['A1']
    tc.value = title
    tc.font = Font(bold=True, size=14, color='1E40AF')
    tc.alignment = Alignment(horizontal='center')
    ws.row_dimensions[1].height = 28

    fill = PatternFill(fill_type='solid', fgColor='1E40AF')
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=2, column=col, value=h)
        cell.font = Font(bold=True, color='FFFFFF')
        cell.fill = fill
        cell.alignment = Alignment(horizontal='center')

    for r, row in enumerate(rows, 3):
        for c, val in enumerate(row, 1):
            ws.cell(row=r, column=c, value=val)

    ws.column_dimensions['A'].width = 18
    ws.column_dimensions['B'].width = 18
    ws.column_dimensions['C'].width = 20
    ws.column_dimensions['D'].width = 20

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


def _generate_template_file(name, fmt):
    body = (
        f'Настоящий документ «{name}» составлен в соответствии с '
        'внутренними регламентами организации и является обязательным '
        'для исполнения всеми сотрудниками соответствующего подразделения.'
    )
    safe_name = name.replace(' ', '_').replace('/', '-').replace('«', '').replace('»', '')
    if fmt == 'docx':
        data = _make_docx(name, body)
        filename = f'{safe_name}.docx'
    elif fmt == 'xlsx':
        data = _make_xlsx(name)
        filename = f'{safe_name}.xlsx'
    else:
        data = _make_docx(name, body)
        filename = f'{safe_name}.docx'
    return ContentFile(data), filename


# ─── Management command ───────────────────────────────────────────────────────

class Command(BaseCommand):
    help = 'Инициализация production БД: пользователи + тестовые данные'

    def handle(self, *args, **options):
        self._ensure_users()
        if Document.objects.count() == 0:
            self.stdout.write('  БД пуста — загружаю тестовые данные...')
            self._seed_templates()
            self._seed_documents()
            self.stdout.write(self.style.SUCCESS('✅ Тестовые данные загружены'))
        else:
            self.stdout.write('  Документы уже есть — пропускаю сидинг.')

    def _ensure_users(self):
        for u in REQUIRED_USERS:
            user, created = User.objects.get_or_create(username=u['username'])
            user.set_password(u['password'])
            user.first_name = u['first_name']
            user.last_name = u['last_name']
            user.email = u['email']
            user.is_superuser = u['is_superuser']
            user.is_staff = u['is_staff']
            user.save()
            profile, _ = UserProfile.objects.get_or_create(user=user)
            profile.role = u['role']
            profile.position = u['position']
            profile.department = u['department']
            profile.save()
            self.stdout.write(
                self.style.SUCCESS(
                    f'  ✅ {"Создан" if created else "OK"}: {u["username"]} ({u["role"]})'
                )
            )

        idx = 1
        for last, first, role, position, dept in EXTRA_USERS:
            username = f'{last.lower()}.{first.lower()}{idx}'
            idx += 1
            user, created = User.objects.get_or_create(username=username)
            if created:
                user.set_password('Test1234!')
                user.first_name = first
                user.last_name = last
                user.email = f'{username}@example.com'
                user.save()
                profile, _ = UserProfile.objects.get_or_create(user=user)
                profile.role = role
                profile.position = position
                profile.department = dept
                profile.save()

    def _seed_templates(self):
        for name, type_, fmt, desc in TEMPLATES_DATA:
            tpl, created = DocumentTemplate.objects.get_or_create(name=name)
            tpl.type = type_
            tpl.file_format = fmt
            tpl.description = desc
            tpl.is_active = True
            tpl.html_template = f'<h1>{name}</h1><p>{{{{ФИО}}}}, {{{{отдел}}}}</p>'
            file_content, filename = _generate_template_file(name, fmt)
            tpl.template_file.save(filename, file_content, save=False)
            tpl.save()
            self.stdout.write(f'  📄 Шаблон: {name} ({fmt})')

    def _seed_documents(self):
        all_users = list(User.objects.all())
        templates = list(DocumentTemplate.objects.all())

        titles = [
            'Приказ о приёме на работу {name}',
            'Приказ об отпуске {name}',
            'Приказ о премировании сотрудников {dept}',
            'Приказ о командировке {name}',
            'Договор № {num} с ООО «{company}»',
            'Акт приёма-передачи № {num}',
            'Служебная записка от {dept}',
            'Докладная записка о закупке оборудования',
            'Письмо в адрес ООО «{company}»',
            'Финансовый отчёт за {period}',
            'Квартальный отчёт {dept} за {period}',
            'Заявление на отпуск {name}',
            'Трудовой договор с {name}',
            'Смета расходов {dept} на {period}',
            'Акт выполненных работ по договору № {num}',
        ]
        companies = ['Альфа', 'Бета', 'Прогресс', 'Инновация', 'Развитие', 'Омега']
        periods = ['январь 2026', 'февраль 2026', 'март 2026', 'I кв. 2026']
        statuses = ['draft', 'sent_for_approval', 'coordination', 'approved',
                    'execution', 'rejected', 'returned', 'archived']
        weights = [0.10, 0.15, 0.15, 0.25, 0.15, 0.05, 0.05, 0.10]

        for _ in range(60):
            creator = random.choice(all_users)
            tpl = random.choice(templates) if templates else None
            dept = getattr(getattr(creator, 'profile', None), 'department', 'Отдел')

            title_fmt = random.choice(titles)
            title = title_fmt.format(
                name=f'{creator.first_name} {creator.last_name}',
                dept=dept,
                num=f'{random.randint(100, 999)}/2026',
                company=random.choice(companies),
                period=random.choice(periods),
            )

            days_ago = random.randint(0, 180)
            created_at = timezone.now() - timedelta(days=days_ago)
            status = random.choices(statuses, weights=weights)[0]
            assigned = random.choice(all_users) if status != 'draft' else None

            doc = Document.objects.create(
                title=title,
                template=tpl,
                status=status,
                created_by=creator,
                assigned_to=assigned,
                content=(
                    f'Содержание документа: {title}.\n\n'
                    f'Составлен {created_at.strftime("%d.%m.%Y")}.'
                ),
                deadline=(
                    (created_at + timedelta(days=random.randint(5, 30))).date()
                    if random.random() > 0.4 else None
                ),
            )
            Document.objects.filter(pk=doc.pk).update(created_at=created_at)

        self.stdout.write('  📝 Создано 60 документов')
